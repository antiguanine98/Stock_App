"""생약표준품 재고 엑셀 파싱 · 연도별 소급 보정 · AI 프롬프트 구성."""

from __future__ import annotations

import html
import json
import re
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Optional, Sequence, Union

import pandas as pd

# 표/식별에 쓰는 기본 메타 (등록일자 제외, 잔고→재고)
META_PREFERRED = [
    "순번",
    "표준품구분",
    "관리번호",
    "한글명",
    "영문명",
    "재고",
    "분양여부",
]

BALANCE_ALIASES = ("재고", "잔고")
STD_TYPE_ALIASES = {
    "대조품": "대조생약",
    "표준생약": "표준생약",
    "지표성분": "지표성분",
    "대조생약": "대조생약",
}
BULK_DECREASE_THRESHOLD = 100
RELIABILITY_LOW_MAX = 2
RELIABILITY_MID_MAX = 4
MANUFACTURE_CANDIDATE_LIMIT = 10
LONG_TERM_LOW_YEARS = 5.0
LONG_TERM_LOW_REL_DROP = 0.05  # 5년+ 구간 상대 감소율 5% 미만 → 저분양
ZERO_STOCK_CATEGORY = "재고 없음(미보유)"
DEPLETION_CATEGORY_ORDER = (
    "1년 이내",
    "2년 이내",
    "3년 이내",
    "4년 이내",
    "5년 이내",
    "5년 초과/안정",
    ZERO_STOCK_CATEGORY,
)
PRICE_COL_KEYWORDS = ("가격",)
ACCELERATION_FORMULA_KO = (
    "분양 가속도 = 최근 3년 연평균 분양량(감소구간만) ÷ 과거 연평균 분양량; "
    "급가속=비율≥2, 증가=비율≥1.25"
)
PRIORITY_FORMULA_KO = (
    "제조우선순위점수 f = 0.40×재고위험도 + 0.25×최근분양속도(정규화) "
    "+ 0.20×분양가속도(정규화) + 0.15×데이터신뢰도\n"
    "  · 재고위험도 = 1 − (예상소진년수/15)  (0~1, 소진년수 없으면 0.05)\n"
    "  · 최근분양속도 정규화 = min(1, 연평균분양량/50)\n"
    "  · 분양가속도 정규화 = min(1, (가속도비율−1)/2)  (비율 없으면 0)\n"
    "  · 데이터신뢰도 = A:1.0 / B:0.75 / C:0.5 / D:0.25"
)
PathLike = Union[str, Path]
PathList = Sequence[PathLike]


def _excel_engine_for_path(path: PathLike) -> str:
    """확장자 기준 pandas 엔진: .xls → xlrd, 그 외(xlsx 등) → openpyxl."""
    ext = Path(str(path)).suffix.lower()
    if ext == ".xls":
        return "xlrd"
    return "openpyxl"


def _is_zip_or_openpyxl_failure(exc: BaseException) -> bool:
    """openpyxl이 zip(.xlsx)이 아닌 파일을 만났을 때 나는 오류 판별."""
    if isinstance(exc, zipfile.BadZipFile):
        return True
    name = type(exc).__name__
    if "BadZipFile" in name or "BadZipfile" in name:
        return True
    msg = str(exc).lower()
    return "not a zip file" in msg or "file is not a zip file" in msg


def read_excel_dataframe(path: PathLike, **kwargs: Any) -> pd.DataFrame:
    """엑셀 DataFrame 로드 — .xls=xlrd, .xlsx=openpyxl, BadZipFile 시 xlrd fallback."""
    path_s = str(path)
    engine = _excel_engine_for_path(path_s)
    try:
        return pd.read_excel(path_s, engine=engine, **kwargs)
    except Exception as e:
        if engine == "openpyxl" and _is_zip_or_openpyxl_failure(e):
            return pd.read_excel(path_s, engine="xlrd", **kwargs)
        raise


def open_excel_file(path: PathLike) -> pd.ExcelFile:
    """다중 시트용 ExcelFile — 엔진 분기 + BadZipFile → xlrd fallback."""
    path_s = str(path)
    engine = _excel_engine_for_path(path_s)
    try:
        return pd.ExcelFile(path_s, engine=engine)
    except Exception as e:
        if engine == "openpyxl" and _is_zip_or_openpyxl_failure(e):
            return pd.ExcelFile(path_s, engine="xlrd")
        raise


def normalize_excel_path(path: Any) -> Optional[str]:
    """엑셀 경로 정규화. QUrl/튜플→str, ~$·0바이트·비엑셀은 None."""
    if path is None:
        return None
    if hasattr(path, "toLocalFile") and callable(getattr(path, "toLocalFile")):
        try:
            path = path.toLocalFile()
        except Exception:
            return None
    if isinstance(path, (tuple, list)):
        path = path[0] if path else None
    if path is None:
        return None
    if not isinstance(path, str):
        path = str(path)
    path = path.strip()
    if not path:
        return None
    name = Path(path).name
    if name.startswith("~$"):
        return None
    if not path.lower().endswith((".xlsx", ".xls")):
        return None
    try:
        p = Path(path)
        if not p.is_file() or p.stat().st_size <= 0:
            return None
    except OSError:
        return None
    try:
        return str(p.resolve())
    except OSError:
        return path


def _is_empty(value: Any) -> bool:
    if value is None:
        return True
    try:
        if pd.isna(value):
            return True
    except (TypeError, ValueError):
        pass
    if isinstance(value, str) and value.strip() == "":
        return True
    return False


def parse_date(value: Any) -> Optional[date]:
    if _is_empty(value):
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if isinstance(value, pd.Timestamp):
        if pd.isna(value):
            return None
        return value.date()
    text = str(value).strip()
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y%m%d"):
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue
    ts = pd.to_datetime(text, errors="coerce")
    if pd.isna(ts):
        return None
    return ts.date()


def parse_qty(value: Any) -> Optional[float]:
    if _is_empty(value):
        return None
    if isinstance(value, str):
        value = value.strip().replace(",", "")
        if not value:
            return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def format_qty_int(value: Any) -> str:
    """표·리포트용 재고 수량 — 소수점 없는 정수 문자열."""
    if _is_empty(value):
        return "-"
    try:
        num = float(value)
    except (TypeError, ValueError):
        return "-"
    if num != num:  # NaN
        return "-"
    return str(int(round(num)))


def to_qty_int(value: Any) -> Optional[int]:
    """재고 수량을 int로 정규화 (None 유지)."""
    if _is_empty(value):
        return None
    try:
        num = float(value)
    except (TypeError, ValueError):
        return None
    if num != num:
        return None
    return int(round(num))


def is_zero_stock(item: "StockItem") -> bool:
    """현재 재고(잔고)가 0 이하인지 (로트/관리번호 단위)."""
    qty = item.current_qty
    if qty is None:
        return False
    return float(qty) <= 0


def build_name_ko_stock_map(items: list["StockItem"]) -> dict[str, dict[str, Any]]:
    """한글명 기준 재고 통합 맵 — 동일 한글명 로트 합산·보유 여부."""
    out: dict[str, dict[str, Any]] = {}
    for it in items:
        key = (it.name_ko or "").strip()
        if not key:
            key = f"__unnamed__:{it.manage_no or it.label}"
        slot = out.setdefault(
            key,
            {
                "name_ko": (it.name_ko or "").strip() or key,
                "total_qty": 0.0,
                "has_stock": False,
                "lot_count": 0,
                "manage_nos": [],
                "labels": [],
            },
        )
        qty = it.current_qty
        q = float(qty) if qty is not None else 0.0
        if q < 0:
            q = 0.0
        slot["total_qty"] += q
        slot["lot_count"] += 1
        if q > 0:
            slot["has_stock"] = True
        if it.manage_no:
            slot["manage_nos"].append(it.manage_no)
        slot["labels"].append(it.label)
    for slot in out.values():
        slot["total_qty"] = round(float(slot["total_qty"]), 6)
    return out


def name_ko_group_has_stock(
    item: "StockItem",
    name_ko_map: dict[str, dict[str, Any]] | None = None,
    items: list["StockItem"] | None = None,
) -> bool:
    """동일 한글명 품목군에 재고가 하나라도 있으면 True."""
    if name_ko_map is None:
        name_ko_map = build_name_ko_stock_map(items or [item])
    key = (item.name_ko or "").strip() or f"__unnamed__:{item.manage_no or item.label}"
    slot = name_ko_map.get(key)
    if slot is None:
        return not is_zero_stock(item)
    return bool(slot.get("has_stock"))


def is_name_level_zero_stock(
    item: "StockItem",
    name_ko_map: dict[str, dict[str, Any]] | None = None,
    items: list["StockItem"] | None = None,
) -> bool:
    """한글명 단위 미보유 — 동일 한글명 전 로트가 재고 0일 때만 True."""
    if name_ko_map is None:
        name_ko_map = build_name_ko_stock_map(items or [item])
    key = (item.name_ko or "").strip() or f"__unnamed__:{item.manage_no or item.label}"
    slot = name_ko_map.get(key)
    if slot is None:
        return is_zero_stock(item)
    return not bool(slot.get("has_stock"))


def inventory_identity_key(item: "StockItem") -> str:
    """공정서 매칭용 동일 품목 키 — 영문명 우선, 없으면 한글명."""
    en = _cell_str(getattr(item, "name_en", None)).strip()
    if en:
        nk = _norm_key(en)
        if nk:
            return f"en:{nk}"
    ko = (item.name_ko or "").strip()
    if ko:
        nk = _norm_key(ko)
        if nk:
            return f"ko:{nk}"
    return f"id:{item.manage_no or item.label or id(item)}"


def build_name_en_inventory_groups(
    items: list["StockItem"],
) -> dict[str, list["StockItem"]]:
    """영문명(없으면 한글명) 동일 품목을 1그룹으로 묶음."""
    groups: dict[str, list[StockItem]] = defaultdict(list)
    for it in items:
        groups[inventory_identity_key(it)].append(it)
    return dict(groups)


def filter_zero_stock_rows_by_name_ko(
    rows: list[dict[str, Any]],
    name_ko_map: dict[str, dict[str, Any]],
) -> list[dict[str, Any]]:
    """재고 없음 목록 — 한글명 그룹에 재고가 있으면 제외, 한글명당 1행."""
    seen: set[str] = set()
    out: list[dict[str, Any]] = []
    for r in rows:
        name = str(r.get("name_ko") or "").strip()
        key = name or str(r.get("manage_no") or r.get("label") or "")
        if not key or key in seen:
            continue
        slot = name_ko_map.get(name) if name else None
        if slot is not None and slot.get("has_stock"):
            continue
        if name and name not in name_ko_map:
            # 맵에 없으면 행의 stock_zero/last_qty로 판정
            if not r.get("stock_zero") and (r.get("last_qty") or 0) not in (0, "0", None, "-"):
                continue
        seen.add(key)
        row = dict(r)
        if slot is not None:
            row["last_qty"] = to_qty_int(slot.get("total_qty"))
            row["lot_count"] = slot.get("lot_count")
            row["group_manage_nos"] = list(slot.get("manage_nos") or [])
        out.append(row)
    return out


def format_date(d: date) -> str:
    """호환용 전체 일자 표기 (내부·대량출고 탐지 등)."""
    return d.strftime("%Y-%m-%d")


def format_year(d: date | int) -> str:
    """UI/차트용 연도(YYYY) 표기."""
    if isinstance(d, int):
        return str(d)
    return str(d.year)


def _cell_str(value: Any) -> str:
    if _is_empty(value):
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def normalize_std_type(raw: Any) -> str:
    """생약(표준생약) → 표준생약, 생약(대조품) → 대조생약 등으로 정제."""
    text = _cell_str(raw)
    if not text:
        return ""
    match = re.search(r"\(([^)]+)\)", text)
    core = match.group(1).strip() if match else text
    core = re.sub(r"^생약\s*", "", core).strip() or core
    return STD_TYPE_ALIASES.get(core, core)


def display_column_name(name: str) -> str:
    """엑셀 원본 컬럼명을 UI 표기명으로 변환 (잔고→재고, 등록일자 제외는 상위에서 처리)."""
    text = str(name).strip()
    if text == "잔고":
        return "재고"
    return text


def discover_timeseries_pairs(columns: list[str]) -> list[tuple[str, str]]:
    """컬럼명에 '변경일자'/'재고량' 키워드가 있는 열을 동적으로 짝지음.

    고정 열 인덱스를 사용하지 않으며, 규격·단위·가격 등 중간 컬럼은 무시한다.
    """
    date_map: dict[int, str] = {}
    qty_map: dict[int, str] = {}

    for col in columns:
        name = str(col).strip()
        if "등록일자" in name:
            continue
        if "변경일자" in name:
            m = re.search(r"변경일자\s*(\d+)", name)
            key = int(m.group(1)) if m else (10_000 + len(date_map))
            date_map[key] = col
        elif "재고량" in name:
            m = re.search(r"재고량\s*(\d+)", name)
            key = int(m.group(1)) if m else (10_000 + len(qty_map))
            qty_map[key] = col

    pairs: list[tuple[str, str]] = []
    for key in sorted(set(date_map) & set(qty_map)):
        pairs.append((date_map[key], qty_map[key]))
    return pairs


def classify_columns(columns: list[str]) -> tuple[list[str], list[tuple[str, str]]]:
    """메타 컬럼과 시계열 쌍을 분리. 등록일자는 메타에서 제외."""
    pairs = discover_timeseries_pairs(columns)
    pair_cols = {c for pair in pairs for c in pair}
    meta_cols: list[str] = []
    for col in columns:
        name = str(col).strip()
        if col in pair_cols:
            continue
        if "등록일자" in name:
            continue
        meta_cols.append(col)

    # 선호 순서 정렬 (잔고/재고 통합)
    ordered: list[str] = []
    used: set[str] = set()
    balance_col = next((c for c in meta_cols if str(c).strip() in BALANCE_ALIASES), None)

    for preferred in META_PREFERRED:
        if preferred == "재고":
            if balance_col and balance_col not in used:
                ordered.append(balance_col)
                used.add(balance_col)
            continue
        hit = next((c for c in meta_cols if str(c).strip() == preferred), None)
        if hit and hit not in used:
            ordered.append(hit)
            used.add(hit)

    for col in meta_cols:
        if col not in used:
            ordered.append(col)
            used.add(col)

    return ordered, pairs


@dataclass
class StockPoint:
    change_date: date
    quantity: float


@dataclass
class StockItem:
    seq: Any = None
    std_type: str = ""
    std_type_raw: Any = None
    manage_no: str = ""
    name_ko: str = ""
    name_en: Any = None
    balance: Any = None
    registered_date: Optional[date] = None
    distributed: Any = None
    unit_price: Optional[float] = None
    extra_meta: dict[str, Any] = field(default_factory=dict)
    raw_points: list[StockPoint] = field(default_factory=list)
    year_end_points: list[StockPoint] = field(default_factory=list)
    corrected_points: list[StockPoint] = field(default_factory=list)
    correction_count: int = 0

    @property
    def label(self) -> str:
        if self.manage_no and self.name_ko:
            return f"{self.manage_no}({self.name_ko})"
        return self.manage_no or self.name_ko or "(미식별)"

    @property
    def first_qty(self) -> Optional[float]:
        if not self.corrected_points:
            return None
        return self.corrected_points[0].quantity

    @property
    def last_qty(self) -> Optional[float]:
        if not self.corrected_points:
            return None
        return self.corrected_points[-1].quantity

    @property
    def current_qty(self) -> Optional[float]:
        """현 재고: 보정 최종값 우선, 없으면 잔고/재고 컬럼."""
        if self.last_qty is not None:
            return self.last_qty
        return parse_qty(self.balance)

    @property
    def stock_value(self) -> Optional[float]:
        """현재 재고 × 가격(원)."""
        qty = self.current_qty
        if qty is None or self.unit_price is None:
            return None
        return float(qty) * float(self.unit_price)

    @property
    def qty_delta(self) -> Optional[float]:
        if self.first_qty is None or self.last_qty is None:
            return None
        return self.last_qty - self.first_qty

    @property
    def has_stock_change(self) -> bool:
        delta = self.qty_delta
        if delta is None:
            return False
        return abs(delta) > 1e-9


def collapse_same_dates(points: list[StockPoint]) -> list[StockPoint]:
    """동일 변경일자에 여러 값이 있으면 마지막 값만 유지 후 날짜 정렬."""
    by_date: dict[date, float] = {}
    for point in points:
        by_date[point.change_date] = point.quantity
    return [StockPoint(d, by_date[d]) for d in sorted(by_date.keys())]


def collapse_to_year_end(points: list[StockPoint]) -> list[StockPoint]:
    """연도별 최종 재고량(해당 연도 마지막 기록일·수량)만 남김."""
    same_day = collapse_same_dates(points)
    year_end: dict[int, StockPoint] = {}
    for point in same_day:
        year_end[point.change_date.year] = point
    return [year_end[y] for y in sorted(year_end.keys())]


def fill_continuous_years(
    points: list[StockPoint],
    *,
    fill_forward: bool = True,
    year_min: int | None = None,
    year_max: int | None = None,
) -> list[StockPoint]:
    """연도 축이 끊기지 않도록 빈 연도를 직전 수량으로 채운다.

    분석(분양속도 등)에는 원본 연도말 포인트를 쓰고, 표/차트 표시에만 사용한다.
    """
    if not points and (year_min is None or year_max is None):
        return []
    by_year = {p.change_date.year: p.quantity for p in points}
    if not by_year and year_min is not None and year_max is not None:
        return []
    y0 = year_min if year_min is not None else min(by_year)
    y1 = year_max if year_max is not None else max(by_year)
    if y1 < y0:
        return []
    result: list[StockPoint] = []
    last: float | None = None
    for y in range(y0, y1 + 1):
        if y in by_year:
            last = by_year[y]
            result.append(StockPoint(date(y, 12, 31), last))
        elif fill_forward and last is not None:
            result.append(StockPoint(date(y, 12, 31), last))
        else:
            result.append(StockPoint(date(y, 12, 31), float("nan")))
    return result


def align_display_series(
    corrected: list[StockPoint],
    original: list[StockPoint],
) -> tuple[list[str], list[float | None], list[float | None]]:
    """보정·원본을 공통 연속 연도축으로 맞춘다. 반환: years, corr_qtys, orig_qtys."""
    years_set = {p.change_date.year for p in corrected} | {
        p.change_date.year for p in original
    }
    if not years_set:
        return [], [], []
    y0, y1 = min(years_set), max(years_set)
    corr_filled = fill_continuous_years(corrected, year_min=y0, year_max=y1)
    orig_map = {p.change_date.year: p.quantity for p in original}
    years = [format_year(p.change_date) for p in corr_filled]
    corr_qtys: list[float | None] = []
    orig_qtys: list[float | None] = []
    last_orig: float | None = None
    for p in corr_filled:
        y = p.change_date.year
        cq = p.quantity
        corr_qtys.append(None if cq != cq else cq)  # NaN → None
        if y in orig_map:
            last_orig = orig_map[y]
            orig_qtys.append(last_orig)
        else:
            orig_qtys.append(last_orig)
    return years, corr_qtys, orig_qtys


def apply_retroactive_correction(points: list[StockPoint]) -> tuple[list[StockPoint], int]:
    """연도별 최종 재고량 기준 소급 보정.

    1) 동일일자 → 마지막 값
    2) 연도별 최종 기록만 추출
    3) 재고가 증가한 연도의 차액만큼 과거 연도 수량을 상향 조정
    """
    year_points = collapse_to_year_end(points)
    if not year_points:
        return [], 0

    quantities = [p.quantity for p in year_points]
    corrected = list(quantities)
    correction_count = 0

    for j in range(1, len(corrected)):
        prev_val = corrected[j - 1]
        curr_val = corrected[j]
        if curr_val > prev_val:
            delta = curr_val - prev_val
            for k in range(j):
                corrected[k] += delta
                correction_count += 1

    result = [
        StockPoint(change_date=p.change_date, quantity=corrected[i])
        for i, p in enumerate(year_points)
    ]
    return result, correction_count


def extract_change_pairs(row: pd.Series, pairs: list[tuple[str, str]]) -> list[StockPoint]:
    """동적 탐색된 변경일자/재고량 쌍에서 유효 포인트만 추출."""
    points: list[StockPoint] = []
    for date_col, qty_col in pairs:
        d = parse_date(row.get(date_col))
        q = parse_qty(row.get(qty_col))
        if d is not None and q is not None:
            points.append(StockPoint(change_date=d, quantity=q))
    return points


def find_balance_value(row: pd.Series, columns: list[str]) -> Any:
    for alias in BALANCE_ALIASES:
        if alias in columns:
            return row.get(alias)
    return None


def find_unit_price(row: pd.Series, columns: list[str]) -> Optional[float]:
    """가격(원) 등 가격 키워드 컬럼에서 단가 추출."""
    for col in columns:
        name = str(col).strip()
        if any(k in name for k in PRICE_COL_KEYWORDS):
            q = parse_qty(row.get(col))
            if q is not None and q >= 0:
                return q
    return None


def _extract_unit_price_from_meta(extra_meta: dict[str, Any]) -> Optional[float]:
    for key, val in extra_meta.items():
        if any(k in str(key) for k in PRICE_COL_KEYWORDS):
            q = parse_qty(val)
            if q is not None and q >= 0:
                return q
    return None


def build_corrected_dataframe(
    items: list[StockItem],
    meta_cols: list[str],
) -> pd.DataFrame:
    """소급 보정 최종 수치만 반영한 표용 DataFrame (연도 YYYY · 연속 연도축)."""
    display_series = [align_display_series(it.corrected_points, it.year_end_points) for it in items]
    max_points = max((len(s[0]) for s in display_series), default=0)
    display_meta = [display_column_name(c) for c in meta_cols]
    records: list[dict[str, Any]] = []

    for it, (years, corr_qtys, _orig_qtys) in zip(items, display_series):
        rec: dict[str, Any] = {}
        for src, disp in zip(meta_cols, display_meta):
            key = str(src).strip()
            if key in BALANCE_ALIASES:
                rec[disp] = it.balance
            elif key == "표준품구분":
                rec[disp] = it.std_type
            elif key == "관리번호":
                rec[disp] = it.manage_no
            elif key == "한글명":
                rec[disp] = it.name_ko
            elif key == "영문명":
                rec[disp] = it.name_en
            elif key == "순번":
                rec[disp] = it.seq
            elif key == "분양여부":
                rec[disp] = it.distributed
            else:
                rec[disp] = it.extra_meta.get(src, "")

        for idx in range(max_points):
            if idx < len(years):
                rec[f"변경일자{idx + 1}"] = years[idx]
                rec[f"재고량{idx + 1}"] = corr_qtys[idx]
            else:
                rec[f"변경일자{idx + 1}"] = None
                rec[f"재고량{idx + 1}"] = None
        records.append(rec)

    columns = list(display_meta)
    for idx in range(max_points):
        columns += [f"변경일자{idx + 1}", f"재고량{idx + 1}"]
    return pd.DataFrame(records, columns=columns)


def _item_merge_key(item: StockItem) -> str:
    if item.manage_no:
        return f"no:{item.manage_no}"
    return f"name:{item.name_ko}"


def merge_stock_items(groups: Sequence[Sequence[StockItem]]) -> list[StockItem]:
    """여러 파일의 품목을 관리번호(없으면 한글명) 기준으로 통합."""
    merged: dict[str, StockItem] = {}
    order: list[str] = []

    for group in groups:
        for it in group:
            key = _item_merge_key(it)
            if key not in merged:
                merged[key] = StockItem(
                    seq=it.seq,
                    std_type=it.std_type,
                    std_type_raw=it.std_type_raw,
                    manage_no=it.manage_no,
                    name_ko=it.name_ko,
                    name_en=it.name_en,
                    balance=it.balance,
                    registered_date=it.registered_date,
                    distributed=it.distributed,
                    unit_price=it.unit_price,
                    extra_meta=dict(it.extra_meta),
                    raw_points=list(it.raw_points),
                    year_end_points=list(it.year_end_points),
                    corrected_points=list(it.corrected_points),
                    correction_count=it.correction_count,
                )
                order.append(key)
                continue

            existing = merged[key]
            existing.raw_points = list(existing.raw_points) + list(it.raw_points)
            existing.year_end_points = collapse_to_year_end(existing.raw_points)
            existing.corrected_points, existing.correction_count = apply_retroactive_correction(
                existing.raw_points
            )
            if it.std_type:
                existing.std_type = it.std_type
                existing.std_type_raw = it.std_type_raw or existing.std_type_raw
            if it.name_ko:
                existing.name_ko = it.name_ko
            if not _is_empty(it.name_en):
                existing.name_en = it.name_en
            if not _is_empty(it.balance):
                existing.balance = it.balance
            if it.unit_price is not None:
                existing.unit_price = it.unit_price
            if not _is_empty(it.distributed):
                existing.distributed = it.distributed
            if it.registered_date and (
                existing.registered_date is None or it.registered_date > existing.registered_date
            ):
                existing.registered_date = it.registered_date
            existing.extra_meta.update({k: v for k, v in it.extra_meta.items() if not _is_empty(v)})
            if existing.seq is None:
                existing.seq = it.seq

    return [merged[k] for k in order]


def _unify_meta_cols(items: list[StockItem]) -> list[str]:
    """통합 품목 목록에서 표시용 메타 컬럼 순서 결정."""
    present = {"관리번호", "한글명"}
    for it in items:
        if it.seq is not None:
            present.add("순번")
        if it.std_type:
            present.add("표준품구분")
        if not _is_empty(it.name_en):
            present.add("영문명")
        if not _is_empty(it.balance):
            present.add("재고")
        if not _is_empty(it.distributed):
            present.add("분양여부")
        for k in it.extra_meta:
            present.add(str(k))
    ordered: list[str] = []
    for pref in META_PREFERRED:
        if pref in present:
            ordered.append(pref)
            present.discard(pref)
    ordered.extend(sorted(present - set(ordered)))
    return ordered


def load_stock_excel(path: str) -> tuple[pd.DataFrame, list[StockItem]]:
    """엑셀을 읽어 연도별 소급 보정된 StockItem 목록과 표용 DataFrame을 반환."""
    items = load_stock_items(path)
    meta_cols = _unify_meta_cols(items)
    return build_corrected_dataframe(items, meta_cols), items


def load_stock_items(path: PathLike) -> list[StockItem]:
    """단일 엑셀에서 StockItem 목록만 로드."""
    path = str(path)
    df = read_excel_dataframe(path)
    df.columns = [str(c).strip() for c in df.columns]

    missing = [c for c in ("관리번호", "한글명") if c not in df.columns]
    if missing:
        raise ValueError(f"필수 컬럼이 없습니다 ({Path(path).name}): {', '.join(missing)}")

    meta_cols, pairs = classify_columns(list(df.columns))
    if not pairs:
        raise ValueError(
            f"'{Path(path).name}'에서 '변경일자'/'재고량' 시계열 컬럼 쌍을 찾을 수 없습니다."
        )

    known = {"순번", "표준품구분", "관리번호", "한글명", "영문명", "재고", "잔고", "분양여부"}
    reg_col = next((c for c in df.columns if "등록일자" in str(c)), None)
    items: list[StockItem] = []

    for _, row in df.iterrows():
        manage_no = _cell_str(row.get("관리번호"))
        name_ko = _cell_str(row.get("한글명"))
        if not manage_no and not name_ko:
            continue

        raw_points = extract_change_pairs(row, pairs)
        year_end = collapse_to_year_end(raw_points)
        corrected, corr_cnt = apply_retroactive_correction(raw_points)
        registered = parse_date(row.get(reg_col)) if reg_col else None

        extra = {
            c: row.get(c)
            for c in meta_cols
            if str(c).strip() not in known
        }

        items.append(
            StockItem(
                seq=row.get("순번"),
                std_type=normalize_std_type(row.get("표준품구분")),
                std_type_raw=row.get("표준품구분"),
                manage_no=manage_no,
                name_ko=name_ko,
                name_en=row.get("영문명"),
                balance=find_balance_value(row, list(df.columns)),
                registered_date=registered,
                distributed=row.get("분양여부"),
                unit_price=find_unit_price(row, list(df.columns)),
                extra_meta=extra,
                raw_points=raw_points,
                year_end_points=year_end,
                corrected_points=corrected,
                correction_count=corr_cnt,
            )
        )

    return items


def process_excel(file_path: PathLike) -> dict[str, Any]:
    """단일 파일 GUI 처리 (다중 파일 API의 래퍼)."""
    return process_excels([file_path])


def process_excels(file_paths: PathList) -> dict[str, Any]:
    """하나 이상의 엑셀을 통합 데이터셋으로 결합해 GUI용 결과를 반환.

    파일별 try-except로 처리해 단일 실패가 전체를 망가뜨리지 않는다.
    실패한 파일은 failed_files에 남기고, 성공 파일만 병합한다.
    """
    paths = [str(p) for p in file_paths]
    if not paths:
        raise ValueError("업로드할 엑셀 파일이 없습니다.")

    groups: list[list[StockItem]] = []
    loaded_paths: list[str] = []
    failed_files: list[dict[str, str]] = []
    for p in paths:
        try:
            groups.append(load_stock_items(p))
            loaded_paths.append(p)
        except Exception as e:
            failed_files.append(
                {
                    "path": p,
                    "name": Path(p).name,
                    "error": str(e),
                }
            )

    if not groups:
        details = "\n".join(
            f"- {f['name']}: {f['error']}" for f in failed_files
        ) or "(상세 없음)"
        raise ValueError(
            "로드에 성공한 엑셀 파일이 없습니다.\n"
            f"실패한 파일:\n{details}"
        )

    items = merge_stock_items(groups) if len(groups) > 1 else list(groups[0])
    meta_cols = _unify_meta_cols(items)
    table_df = build_corrected_dataframe(items, meta_cols)

    display_series = [
        align_display_series(it.corrected_points, it.year_end_points) for it in items
    ]
    max_pairs = max((len(s[0]) for s in display_series), default=0)
    total_corrections = sum(it.correction_count for it in items)
    categories = sorted({it.std_type for it in items if it.std_type})

    ui_items: list[dict[str, Any]] = []
    for idx, (it, (years, corr_qtys, orig_qtys)) in enumerate(zip(items, display_series)):
        meta_vals = []
        for col in meta_cols:
            if col == "재고":
                meta_vals.append(_cell_str(it.balance))
            elif col == "표준품구분":
                meta_vals.append(it.std_type)
            elif col == "관리번호":
                meta_vals.append(it.manage_no)
            elif col == "한글명":
                meta_vals.append(it.name_ko)
            elif col == "영문명":
                meta_vals.append(_cell_str(it.name_en))
            elif col == "순번":
                meta_vals.append(_cell_str(it.seq))
            elif col == "분양여부":
                meta_vals.append(_cell_str(it.distributed))
            else:
                meta_vals.append(_cell_str(it.extra_meta.get(col, "")))

        ui_items.append(
            {
                "row_index": idx,
                "meta": meta_vals,
                "mgmt_no": it.manage_no,
                "korean_name": it.name_ko,
                "std_type": it.std_type,
                "label": it.label,
                "dates": years,
                "original_dates": years,
                "original": orig_qtys,
                "corrected": corr_qtys,
                "has_change": it.has_stock_change,
                "delta": it.qty_delta if it.qty_delta is not None else 0.0,
                "first_qty": it.first_qty,
                "last_qty": it.last_qty,
                "correction_cells": it.correction_count,
                "stock_item": it,
            }
        )

    names = [Path(p).name for p in loaded_paths]
    return {
        "file_path": loaded_paths[0],
        "file_paths": loaded_paths,
        "file_name": names[0] if len(names) == 1 else f"{len(names)}개 파일",
        "file_names": names,
        "meta_cols": meta_cols,
        "mgmt_idx": meta_cols.index("관리번호") if "관리번호" in meta_cols else 0,
        "items": ui_items,
        "stock_items": items,
        "table_df": table_df,
        "correction_count": total_corrections,
        "row_count": len(items),
        "max_pair_count": max_pairs,
        "categories": categories,
        "failed_files": failed_files,
    }


def items_for_ai_analysis(items: list[StockItem]) -> list[StockItem]:
    return [it for it in items if it.has_stock_change]


def detect_bulk_decrease_dates(items: list[StockItem], threshold: int = BULK_DECREASE_THRESHOLD) -> list[str]:
    """동일 날짜에 threshold개 이상 품목이 감소한 날짜(연구과제 대량출고) 탐지."""
    decrease_counter: Counter[str] = Counter()
    for it in items:
        pts = it.corrected_points
        for i in range(1, len(pts)):
            if pts[i].quantity < pts[i - 1].quantity - 1e-9:
                decrease_counter[format_date(pts[i].change_date)] += 1
    return sorted(d for d, n in decrease_counter.items() if n >= threshold)


def decrease_only_rate_stats(points: list[StockPoint]) -> dict[str, Any]:
    """증가 구간(전수조사/추가제조/반납 등)을 제외한 감소 구간만으로 분양속도 산출."""
    total_drop = 0.0
    total_days = 0
    increase_segments = 0
    decrease_segments = 0
    for i in range(1, len(points)):
        delta = points[i - 1].quantity - points[i].quantity
        days = (points[i].change_date - points[i - 1].change_date).days
        if days <= 0:
            continue
        if delta > 1e-9:
            total_drop += delta
            total_days += days
            decrease_segments += 1
        elif delta < -1e-9:
            increase_segments += 1
    if total_days <= 0 or total_drop <= 1e-9:
        return {
            "annual_rate": None,
            "total_drop": total_drop,
            "decrease_years": 0.0,
            "increase_segments": increase_segments,
            "decrease_segments": decrease_segments,
        }
    years = total_days / 365.25
    return {
        "annual_rate": total_drop / years,
        "total_drop": total_drop,
        "decrease_years": years,
        "increase_segments": increase_segments,
        "decrease_segments": decrease_segments,
    }


def _recent_vs_past_decrease_rates(
    points: list[StockPoint],
    recent_years: float = 3.0,
) -> tuple[Optional[float], Optional[float]]:
    """최근 N년 vs 과거 감소 구간 연평균 속도.

    Returns:
        (past_annual_rate, recent_annual_rate)
        cutoff = last_date - recent_years 기준으로 구간을 나눔.
    """
    if len(points) < 2:
        return None, None
    last_date = points[-1].change_date
    cutoff = last_date - timedelta(days=int(round(recent_years * 365.25)))

    split_idx = None
    for i, p in enumerate(points):
        if p.change_date >= cutoff:
            split_idx = i
            break

    if split_idx is None:
        # 전 구간이 과거(최근 구간 없음)
        past = decrease_only_rate_stats(points)["annual_rate"]
        return past, None
    if split_idx == 0:
        # 전 구간이 최근(과거 구간 없음) — 경계 직전 포인트가 없으면 과거 None
        recent = decrease_only_rate_stats(points)["annual_rate"]
        return None, recent

    # 과거: cutoff 이전(+경계점), 최근: 경계 직전 포인트부터(연속 구간 포함)
    past_pts = points[: split_idx + 1]
    recent_pts = points[split_idx - 1 :]
    past = decrease_only_rate_stats(past_pts)["annual_rate"] if len(past_pts) >= 2 else None
    recent = decrease_only_rate_stats(recent_pts)["annual_rate"] if len(recent_pts) >= 2 else None
    return past, recent


def _half_period_decrease_rate(points: list[StockPoint]) -> tuple[Optional[float], Optional[float]]:
    """호환용: 전반/후반 분할. 신규 로직은 `_recent_vs_past_decrease_rates` 사용."""
    if len(points) < 3:
        return None, None
    mid = len(points) // 2
    early = decrease_only_rate_stats(points[: mid + 1])["annual_rate"]
    late = decrease_only_rate_stats(points[mid:])["annual_rate"]
    return early, late


def acceleration_from_rates(
    early_rate: Optional[float],
    late_rate: Optional[float],
) -> dict[str, Any]:
    """분양 가속도: 급가속 / 증가 / 안정 / 감소.

    비율 = 최근(late) / 과거(early). 급가속=비율≥2, 증가=비율≥1.25,
    감소=비율≤0.75, 그 외 안정. is_surge는 급가속만 True.
    """
    if early_rate is None and late_rate is None:
        return {"label": "안정", "ratio": None, "is_surge": False}
    if early_rate is None or early_rate <= 1e-9:
        if late_rate is not None and late_rate > 1e-9:
            return {"label": "급가속", "ratio": 3.0, "is_surge": True}
        return {"label": "안정", "ratio": None, "is_surge": False}
    if late_rate is None or late_rate <= 1e-9:
        return {"label": "감소", "ratio": 0.0, "is_surge": False}

    ratio = float(late_rate) / float(early_rate)
    if ratio >= 2.0:
        label = "급가속"
    elif ratio >= 1.25:
        label = "증가"
    elif ratio <= 0.75:
        label = "감소"
    else:
        label = "안정"
    return {"label": label, "ratio": ratio, "is_surge": label == "급가속"}


def compute_data_reliability(item: StockItem) -> dict[str, Any]:
    """수집 횟수 + 관측 간격 기반 신뢰도 등급 A~D."""
    pts = collapse_same_dates(item.raw_points)
    n = len(pts)
    avg_gap_years: Optional[float] = None
    if n >= 2:
        gaps = [
            (pts[i].change_date - pts[i - 1].change_date).days / 365.25
            for i in range(1, n)
        ]
        avg_gap_years = sum(gaps) / len(gaps) if gaps else None

    if n >= 6:
        grade, score = "A", 1.0
    elif n >= 4:
        grade, score = "B", 0.75
    elif n >= 3:
        grade, score = "C", 0.5
    else:
        grade, score = "D", 0.25

    if avg_gap_years is not None and avg_gap_years > 3.5 and grade in ("A", "B"):
        grade = "B" if grade == "A" else "C"
        score = 0.75 if grade == "B" else 0.5
    if avg_gap_years is not None and avg_gap_years > 5.0 and grade == "C":
        grade, score = "D", 0.25

    labels = {
        "A": "신뢰도 A (높음)",
        "B": "신뢰도 B (양호)",
        "C": "신뢰도 C (보통)",
        "D": "신뢰도 D (낮음·데이터 부족)",
    }
    return {
        "count": n,
        "grade": grade,
        "level": grade,
        "label": labels[grade],
        "score": float(score),
        "avg_gap_years": round(avg_gap_years, 3) if avg_gap_years is not None else None,
    }


def depletion_bucket(years_left: Optional[float], *, stock_zero: bool = False) -> str:
    """소진 예상 기간 카테고리 (1~5년 연 단위 분할)."""
    if stock_zero:
        return ZERO_STOCK_CATEGORY
    if years_left is None:
        return "5년 초과/안정"
    y = float(years_left)
    if y <= 1:
        return "1년 이내"
    if y <= 2:
        return "2년 이내"
    if y <= 3:
        return "3년 이내"
    if y <= 4:
        return "4년 이내"
    if y <= 5:
        return "5년 이내"
    return "5년 초과/안정"


def format_deplete_ym(
    years_left: float,
    *,
    reference_date: date | None = None,
) -> str:
    """분석 기준일(오늘)부터 잔여수명만큼 더한 예상 소진 시점 — YYYY년 MM월."""
    ref = reference_date or date.today()
    days = max(0, int(round(float(years_left) * 365.25)))
    target = ref + timedelta(days=days)
    today = date.today()
    if target < today:
        target = today
    return f"{target.year}년 {target.month:02d}월"


def _priority_components(
    *,
    years_left: Optional[float],
    annual_rate: Optional[float],
    acceleration_ratio: Optional[float],
    reliability_score: float,
) -> dict[str, float]:
    """제조우선순위 구성요소: risk / speed_n / accel_n / rel / score."""
    if years_left is None:
        risk = 0.05
    elif years_left <= 0:
        risk = 1.0
    else:
        risk = max(0.0, min(1.0, 1.0 - (years_left / 15.0)))

    speed_n = 0.0
    if annual_rate is not None and annual_rate > 0:
        speed_n = max(0.0, min(1.0, float(annual_rate) / 50.0))

    accel_n = 0.0
    if acceleration_ratio is not None and acceleration_ratio > 0:
        accel_n = max(0.0, min(1.0, (float(acceleration_ratio) - 1.0) / 2.0))

    rel = max(0.0, min(1.0, float(reliability_score)))
    score = round(0.40 * risk + 0.25 * speed_n + 0.20 * accel_n + 0.15 * rel, 4)
    return {
        "risk": round(risk, 4),
        "speed_n": round(speed_n, 4),
        "accel_n": round(accel_n, 4),
        "rel": round(rel, 4),
        "score": score,
    }


def manufacturing_priority_score(
    *,
    years_left: Optional[float],
    annual_rate: Optional[float],
    acceleration_ratio: Optional[float],
    reliability_score: float,
) -> float:
    """제조우선순위점수 = f(재고위험도, 최근 분양속도, 분양 가속도, 데이터 신뢰도)."""
    return _priority_components(
        years_left=years_left,
        annual_rate=annual_rate,
        acceleration_ratio=acceleration_ratio,
        reliability_score=reliability_score,
    )["score"]


def is_long_term_low_distribution(item: StockItem, annual_rate: Optional[float]) -> bool:
    """최근 5년 이상 감소가 거의 없는 장기 저분양/과다재고."""
    pts = item.corrected_points
    if len(pts) < 2 or item.first_qty is None or item.last_qty is None:
        return False
    span = (pts[-1].change_date - pts[0].change_date).days / 365.25
    if span < LONG_TERM_LOW_YEARS:
        return False
    if abs(item.first_qty) > 1e-12:
        rel_drop = (item.first_qty - item.last_qty) / item.first_qty
        if rel_drop < LONG_TERM_LOW_REL_DROP:
            return True
    if annual_rate is None or annual_rate < 1.0:
        return True
    return False


def estimate_depletion(item: StockItem) -> dict[str, Any]:
    """분양 속도(증가구간 제외)·가속도·신뢰도·우선순위 산출."""
    if item.unit_price is None:
        item.unit_price = _extract_unit_price_from_meta(item.extra_meta)

    reliability = compute_data_reliability(item)
    pts = item.corrected_points
    empty = {
        "speed": "데이터부족",
        "annual_rate": None,
        "years_left": None,
        "deplete_within_2y": False,
        "deplete_within_5y": False,
        "recent_surge": False,
        "rate_change_ratio": None,
        "acceleration": "안정",
        "acceleration_ratio": None,
        "early_rate": None,  # past
        "late_rate": None,  # recent
        "past_rate": None,
        "recent_rate": None,
        "long_term_low": False,
        "increase_segments_excluded": 0,
        "deplete_ym": None,
        "depletion_category": "5년 초과/안정",
        "stock_zero": False,
        "reliability": reliability,
        "stock_risk": 0.0,
        "priority_score": manufacturing_priority_score(
            years_left=None,
            annual_rate=None,
            acceleration_ratio=None,
            reliability_score=reliability["score"],
        ),
        "unit_price": item.unit_price,
        "stock_value": item.stock_value,
        "as_of_year": date.today().year,
    }

    # 현재 재고 0 → 소진 위험군 제외, 재고 없음(미보유)으로 분리
    if is_zero_stock(item):
        empty.update(
            {
                "speed": "해당없음",
                "stock_zero": True,
                "depletion_category": ZERO_STOCK_CATEGORY,
                "deplete_within_2y": False,
                "deplete_within_5y": False,
                "deplete_ym": None,
                "years_left": None,
                "stock_risk": 0.0,
                "priority_score": 0.0,
                "acceleration": "해당없음",
                "long_term_low": False,
            }
        )
        return empty

    if len(pts) < 2 or item.first_qty is None or item.last_qty is None:
        return empty

    dec = decrease_only_rate_stats(pts)
    annual_rate = dec["annual_rate"]
    # early/late = past/recent (최근 3년 vs 과거)
    early_rate, late_rate = _recent_vs_past_decrease_rates(pts, recent_years=3.0)
    accel = acceleration_from_rates(early_rate, late_rate)
    rate_change_ratio = accel["ratio"]
    recent_surge = bool(accel["label"] == "급가속")
    # 소진 잔여수명: 최종 잔고 ÷ 최근 연평균 분양속도 (없으면 전체 감소구간 평균)
    recent_annual_rate = late_rate if late_rate is not None and late_rate > 1e-9 else annual_rate
    analysis_ref = date.today()

    if recent_annual_rate is None or recent_annual_rate <= 1e-9:
        speed = "느림"
        years_left = None
        deplete_5 = False
        deplete_2 = False
        deplete_ym = None
    else:
        years_left = item.last_qty / recent_annual_rate if item.last_qty > 0 else 0.0
        if recent_annual_rate >= 40:
            speed = "빠름"
        elif recent_annual_rate >= 10:
            speed = "보통"
        else:
            speed = "느림"
        deplete_5 = years_left is not None and years_left <= 5.0
        deplete_2 = years_left is not None and years_left <= 2.0
        deplete_ym = format_deplete_ym(float(years_left), reference_date=analysis_ref)

    category = depletion_bucket(years_left, stock_zero=False)
    long_term_low = is_long_term_low_distribution(item, annual_rate)
    risk = risk_grade_from_years(years_left)
    comps = _priority_components(
        years_left=years_left,
        annual_rate=annual_rate,
        acceleration_ratio=accel["ratio"],
        reliability_score=reliability["score"],
    )
    priority = comps["score"]
    stock_risk = comps["risk"]

    return {
        "speed": speed,
        "annual_rate": annual_rate,
        "years_left": years_left,
        "deplete_within_2y": deplete_2,
        "deplete_within_5y": deplete_5,
        "recent_surge": recent_surge,
        "rate_change_ratio": rate_change_ratio,
        "acceleration": accel["label"],
        "acceleration_ratio": accel["ratio"],
        "early_rate": early_rate,
        "late_rate": late_rate,
        "past_rate": early_rate,
        "recent_rate": late_rate,
        "long_term_low": long_term_low,
        "increase_segments_excluded": dec["increase_segments"],
        "deplete_ym": deplete_ym,
        "depletion_category": category,
        "stock_zero": False,
        "reliability": reliability,
        "stock_risk": stock_risk,
        "speed_n": comps["speed_n"],
        "accel_n": comps["accel_n"],
        "priority_score": priority,
        "unit_price": item.unit_price,
        "stock_value": item.stock_value,
        "as_of_year": analysis_ref.year,
        "analysis_reference_date": analysis_ref.isoformat(),
        "recent_annual_rate": recent_annual_rate,
        "risk_grade": risk,
    }


def _stock_item_stats_key(item: StockItem) -> str:
    """estimate_depletion 캐시 키 (관리번호 우선)."""
    if item.manage_no:
        return item.manage_no
    return f"@{id(item)}:{item.label}"


def get_depletion_stats(
    item: StockItem,
    stats_cache: dict[str, dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """품목별 estimate_depletion — stats_cache가 있으면 동일 품목 재계산 생략."""
    if stats_cache is None:
        return estimate_depletion(item)
    key = _stock_item_stats_key(item)
    cached = stats_cache.get(key)
    if cached is not None:
        return cached
    stats = estimate_depletion(item)
    stats_cache[key] = stats
    return stats


def _catalog_item_row(it: StockItem, stats: dict[str, Any]) -> dict[str, Any]:
    """챗봇 전수 목록용 품목 행(한글명·관리번호·세부 수치)."""
    rel = stats.get("reliability") or {}
    grade = rel.get("grade") if isinstance(rel, dict) else None
    years_left = stats.get("years_left")
    stock_zero = bool(stats.get("stock_zero")) or is_zero_stock(it)
    risk = "재고없음" if stock_zero else risk_grade_from_years(years_left)
    return {
        "label": it.label,
        "name_ko": it.name_ko,
        "manage_no": it.manage_no,
        "std_type": it.std_type,
        "last_qty": to_qty_int(it.last_qty if it.last_qty is not None else it.current_qty),
        "annual_rate": stats.get("annual_rate"),
        "years_left": years_left,
        "deplete_ym": stats.get("deplete_ym"),
        "depletion_category": stats.get("depletion_category"),
        "risk_grade": risk,
        "acceleration": stats.get("acceleration"),
        "acceleration_ratio": stats.get("acceleration_ratio"),
        "priority_score": stats.get("priority_score"),
        "stock_risk": stats.get("stock_risk"),
        "reliability_grade": grade,
        "reliability": rel.get("label") if isinstance(rel, dict) else rel,
        "stock_value": stats.get("stock_value"),
        "stock_zero": stock_zero,
        "deplete_within_2y": bool(stats.get("deplete_within_2y")),
        "deplete_within_5y": bool(stats.get("deplete_within_5y")),
    }


def risk_grade_from_years(years_left: Optional[float]) -> str:
    """소진 잔여년수 기반 위험등급: 위험/경계/주의/안정 (재고 0 제외)."""
    if years_left is None:
        return "안정"
    if years_left <= 2:
        return "위험"
    if years_left <= 5:
        return "경계"
    if years_left <= 10:
        return "주의"
    return "안정"


def _display_risk_grade(row: dict[str, Any]) -> str:
    grade = row.get("risk_grade")
    if grade:
        return str(grade)
    if row.get("stock_zero"):
        return "재고없음"
    return risk_grade_from_years(row.get("years_left"))


def group_by_depletion_category(
    items: list[StockItem],
    stats_cache: dict[str, dict[str, Any]] | None = None,
    name_ko_map: dict[str, dict[str, Any]] | None = None,
) -> dict[str, list[str]]:
    """소진 기간 카테고리별 품목 라벨 목록(리포트 요약용).

    '재고 없음(미보유)'은 한글명 단위 — 동일 한글명에 재고가 있으면 제외.
    """
    name_map = name_ko_map or build_name_ko_stock_map(items)
    buckets: dict[str, list[str]] = {k: [] for k in DEPLETION_CATEGORY_ORDER}
    seen_zero_names: set[str] = set()
    for it in items_for_ai_analysis(items):
        stats = get_depletion_stats(it, stats_cache)
        key = stats["depletion_category"]
        if key == ZERO_STOCK_CATEGORY:
            if not is_name_level_zero_stock(it, name_map):
                continue
            nk = (it.name_ko or "").strip() or it.label
            if nk in seen_zero_names:
                continue
            seen_zero_names.add(nk)
        if key not in buckets:
            buckets[key] = []
        qty = format_qty_int(it.last_qty if it.last_qty is not None else it.current_qty)
        buckets[key].append(f"{it.label} [재고:{qty}]")
    return buckets


def group_by_depletion_category_items(
    items: list[StockItem],
    stats_cache: dict[str, dict[str, Any]] | None = None,
    name_ko_map: dict[str, dict[str, Any]] | None = None,
) -> dict[str, list[dict[str, Any]]]:
    """소진 예상 구간별 전수 품목 리스트(챗봇·리포트 표용).

    '재고 없음(미보유)'은 한글명 단위 — 동일 한글명에 재고가 하나라도 있으면 제외.
    """
    name_map = name_ko_map or build_name_ko_stock_map(items)
    buckets: dict[str, list[dict[str, Any]]] = {k: [] for k in DEPLETION_CATEGORY_ORDER}
    for it in items_for_ai_analysis(items):
        stats = get_depletion_stats(it, stats_cache)
        key = stats["depletion_category"]
        if key == ZERO_STOCK_CATEGORY:
            if not is_name_level_zero_stock(it, name_map):
                # 로트 재고 0이어도 동일 한글명에 잔여 재고가 있으면 미보유에서 제외
                continue
        if key not in buckets:
            buckets[key] = []
        buckets[key].append(_catalog_item_row(it, stats))
    # 미보유 목록은 한글명 기준 고유화
    if buckets.get(ZERO_STOCK_CATEGORY):
        buckets[ZERO_STOCK_CATEGORY] = filter_zero_stock_rows_by_name_ko(
            buckets[ZERO_STOCK_CATEGORY], name_map
        )
    return buckets


def group_by_risk_grade(
    items: list[StockItem],
    stats_cache: dict[str, dict[str, Any]] | None = None,
) -> dict[str, list[dict[str, Any]]]:
    """위험등급(위험/경계/주의/안정)별 전수 — 재고 0은 제외."""
    order = ["위험", "경계", "주의", "안정"]
    buckets: dict[str, list[dict[str, Any]]] = {k: [] for k in order}
    for it in items_for_ai_analysis(items):
        stats = get_depletion_stats(it, stats_cache)
        if stats.get("stock_zero") or is_zero_stock(it):
            continue
        row = _catalog_item_row(it, stats)
        grade = row["risk_grade"]
        if grade not in buckets:
            continue
        buckets[grade].append(row)
    for g in order:
        buckets[g].sort(
            key=lambda r: (
                float(r["years_left"]) if isinstance(r.get("years_left"), (int, float)) else 999.0,
                -float(r["priority_score"] or 0),
            )
        )
    return buckets


def select_manufacture_candidates(
    items: list[StockItem],
    limit_per_type: int = MANUFACTURE_CANDIDATE_LIMIT,
    stats_cache: dict[str, dict[str, Any]] | None = None,
) -> dict[str, list[dict[str, Any]]]:
    """차년도 제조검토대상: 표준생약·지표성분 각각 제조우선순위점수 상위 N건(기본 10).

    5년 이내 소진 여부는 하드 필터가 아니라 결과에 표기만 한다.
    해당 유형 품목이 N건 미만이면 전량을 반환한다.
    """
    result: dict[str, list[dict[str, Any]]] = {"표준생약": [], "지표성분": []}
    scored: list[tuple[float, StockItem, dict[str, Any]]] = []
    for it in items:
        if it.std_type not in result:
            continue
        if is_zero_stock(it):
            continue
        stats = get_depletion_stats(it, stats_cache)
        if stats.get("stock_zero"):
            continue
        scored.append((float(stats["priority_score"]), it, stats))

    scored.sort(key=lambda x: x[0], reverse=True)
    for score, it, stats in scored:
        bucket = result[it.std_type]
        if len(bucket) >= limit_per_type:
            continue
        rel = stats["reliability"]
        rel_score = float(rel.get("score", 0.25)) if isinstance(rel, dict) else 0.25
        comps = _priority_components(
            years_left=stats["years_left"],
            annual_rate=stats["annual_rate"],
            acceleration_ratio=stats["acceleration_ratio"],
            reliability_score=rel_score,
        )
        bucket.append(
            {
                "label": it.label,
                "manage_no": it.manage_no,
                "name_ko": it.name_ko,
                "std_type": it.std_type,
                "priority_score": score,
                "stock_risk": comps["risk"],
                "speed_n": comps["speed_n"],
                "accel_n": comps["accel_n"],
                "reliability_score": comps["rel"],
                "reliability_grade": rel.get("grade") if isinstance(rel, dict) else None,
                "reliability": rel["label"] if isinstance(rel, dict) else str(rel),
                "early_rate": stats.get("early_rate"),
                "late_rate": stats.get("late_rate"),
                "past_rate": stats.get("past_rate", stats.get("early_rate")),
                "recent_rate": stats.get("recent_rate", stats.get("late_rate")),
                "acceleration_ratio": stats["acceleration_ratio"],
                "annual_rate": stats["annual_rate"],
                "years_left": stats["years_left"],
                "deplete_ym": stats["deplete_ym"],
                "depletion_category": stats["depletion_category"],
                "deplete_within_5y": bool(stats.get("deplete_within_5y")),
                "last_qty": to_qty_int(it.last_qty),
                "acceleration": stats["acceleration"],
                "risk_grade": risk_grade_from_years(stats.get("years_left")),
            }
        )
    return result


def select_monitoring_targets(
    items: list[StockItem],
    limit: int = 30,
    stats_cache: dict[str, dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    """분양 가속도 급증(급가속/증가) 모니터링 대상.

    급가속은 전량 반환(limit 미적용). 증가 항목은 급가속 뒤에 이어서 포함
    (limit이 있으면 증가 구간에만 상한 적용, 기본은 사실상 무제한).
    """
    surge_rows: list[dict[str, Any]] = []
    increase_rows: list[dict[str, Any]] = []
    for it in items:
        if is_zero_stock(it):
            continue
        stats = get_depletion_stats(it, stats_cache)
        if stats.get("stock_zero"):
            continue
        accel = stats.get("acceleration")
        if accel not in ("급가속", "증가"):
            continue
        row = {
            "label": it.label,
            "name_ko": it.name_ko,
            "manage_no": it.manage_no,
            "std_type": it.std_type,
            "last_qty": it.last_qty,
            "acceleration": stats["acceleration"],
            "acceleration_ratio": stats["acceleration_ratio"],
            "rate_change_ratio": stats["rate_change_ratio"],
            "annual_rate": stats["annual_rate"],
            "years_left": stats["years_left"],
            "deplete_ym": stats["deplete_ym"],
            "depletion_category": stats["depletion_category"],
            "risk_grade": risk_grade_from_years(stats.get("years_left")),
            "reliability": stats["reliability"]["label"],
            "reliability_grade": stats["reliability"].get("grade"),
            "priority_score": stats["priority_score"],
            "stock_risk": stats.get("stock_risk"),
            "stock_value": stats.get("stock_value"),
            "indicators": (
                f"가속도={stats['acceleration']}"
                + (
                    f", 최근/과거비={stats['acceleration_ratio']:.2f}"
                    if stats["acceleration_ratio"] is not None
                    else ""
                )
            ),
        }
        if accel == "급가속":
            surge_rows.append(row)
        else:
            increase_rows.append(row)

    def _sort_key(r: dict[str, Any]) -> tuple[float, float]:
        return (float(r["acceleration_ratio"] or 0), float(r["priority_score"] or 0))

    surge_rows.sort(key=_sort_key, reverse=True)
    increase_rows.sort(key=_sort_key, reverse=True)
    # 급가속 전량 + 증가 전량 (limit은 API 호환용, 급가속 하드캡 없음)
    _ = limit
    return surge_rows + increase_rows


def select_long_term_low_items(
    items: list[StockItem],
    limit: int = 40,
    stats_cache: dict[str, dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    """장기 저분양/과다재고 — 차기 제조 시 수량 하향 조정 권고."""
    rows: list[dict[str, Any]] = []
    for it in items:
        stats = get_depletion_stats(it, stats_cache)
        if not stats.get("long_term_low"):
            continue
        rows.append(
            {
                "label": it.label,
                "manage_no": it.manage_no,
                "std_type": it.std_type,
                "annual_rate": stats["annual_rate"],
                "years_left": stats["years_left"],
                "stock_value": stats.get("stock_value"),
                "reliability": stats["reliability"]["label"],
                "recommendation": "차기 제조 시 수량 하향 조정 검토",
            }
        )
    rows.sort(key=lambda r: float(r["stock_value"] or 0), reverse=True)
    return rows[:limit]


def compute_inventory_valuation(items: list[StockItem]) -> dict[str, Any]:
    """현재 재고 × 가격(원) 환산액 집계."""
    by_type: dict[str, float] = defaultdict(float)
    rows: list[dict[str, Any]] = []
    missing_price = 0
    for it in items:
        if it.unit_price is None:
            it.unit_price = _extract_unit_price_from_meta(it.extra_meta)
        val = it.stock_value
        if val is None:
            missing_price += 1
            continue
        tkey = it.std_type or "미분류"
        by_type[tkey] += val
        rows.append(
            {
                "label": it.label,
                "manage_no": it.manage_no,
                "std_type": tkey,
                "qty": it.current_qty,
                "unit_price": it.unit_price,
                "stock_value": val,
            }
        )
    rows.sort(key=lambda r: float(r["stock_value"]), reverse=True)
    total = sum(by_type.values())
    type_order = ["표준생약", "지표성분", "대조생약", "미분류"]
    by_type_ordered: dict[str, float] = {}
    for k in type_order:
        if k in by_type:
            by_type_ordered[k] = by_type[k]
    for k, v in by_type.items():
        if k not in by_type_ordered:
            by_type_ordered[k] = v
    return {
        "total_value": total,
        "by_type": dict(by_type_ordered),
        "top20": rows[:20],
        "priced_count": len(rows),
        "missing_price_count": missing_price,
    }


def _fmt_money(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{value:,.0f}원"


def build_kpi_dashboard(items: list[StockItem], flags: dict[str, Any] | None = None) -> dict[str, Any]:
    """핵심 KPI + 자동 종합 의견(5줄 내외)."""
    flags = flags or {}
    valuation = flags.get("valuation") or compute_inventory_valuation(items)
    by_code = flags.get("by_code") or {}
    category_items = flags.get("depletion_category_items") or {}
    stats_cache = flags.get("_depletion_stats_cache")

    def _stats_for(it: StockItem) -> dict[str, Any]:
        cached = by_code.get(it.manage_no) if it.manage_no else None
        if cached is not None:
            return cached
        return get_depletion_stats(it, stats_cache)

    managed = len(items)
    # 소진 예상: 1년 내 / 1~3년 / 3~5년 (상호 배타 구간, 재고 0 제외)
    deplete_1y = 0
    deplete_1_3y = 0
    deplete_3_5y = 0
    if category_items:
        deplete_1y = len(category_items.get("1년 이내") or [])
        deplete_1_3y = len(category_items.get("2년 이내") or []) + len(
            category_items.get("3년 이내") or []
        )
        deplete_3_5y = len(category_items.get("4년 이내") or []) + len(
            category_items.get("5년 이내") or []
        )
    else:
        for it in items:
            stats = _stats_for(it)
            if stats.get("stock_zero") or is_zero_stock(it):
                continue
            yl = stats.get("years_left")
            if not isinstance(yl, (int, float)):
                continue
            if yl <= 1:
                deplete_1y += 1
            elif yl <= 3:
                deplete_1_3y += 1
            elif yl <= 5:
                deplete_3_5y += 1

    accel_n = 0
    low_n = 0
    grade_ab = 0
    for it in items:
        stats = _stats_for(it)
        if stats.get("stock_zero") or is_zero_stock(it):
            continue
        if stats.get("acceleration") in ("급가속", "증가"):
            accel_n += 1
        if stats.get("long_term_low"):
            low_n += 1
        rel = stats.get("reliability") or {}
        if isinstance(rel, dict) and rel.get("grade") in ("A", "B"):
            grade_ab += 1

    zero_n = sum(
        1
        for slot in (flags.get("name_ko_stock_map") or build_name_ko_stock_map(items)).values()
        if not slot.get("has_stock")
    )
    deplete_5y_total = deplete_1y + deplete_1_3y + deplete_3_5y

    mfg = flags.get("manufacture_candidates")
    if mfg is None:
        mfg = select_manufacture_candidates(items)
    mfg_n = len(mfg.get("표준생약") or []) + len(mfg.get("지표성분") or [])
    total_value = float(valuation.get("total_value") or 0)
    by_type = valuation.get("by_type") or {}

    kpis = [
        {"key": "managed", "label": "대상품목 수", "value": managed, "display": f"{managed}종"},
        {"key": "zero_stock", "label": "재고 없음(미보유)", "value": zero_n, "display": f"{zero_n}종"},
        {"key": "deplete_1y", "label": "1년 내 소진예상", "value": deplete_1y, "display": f"{deplete_1y}종"},
        {"key": "deplete_1_3y", "label": "1~3년 소진예상", "value": deplete_1_3y, "display": f"{deplete_1_3y}종"},
        {"key": "deplete_3_5y", "label": "3~5년 소진예상", "value": deplete_3_5y, "display": f"{deplete_3_5y}종"},
        {"key": "manufacture", "label": "제조 우선검토 수", "value": mfg_n, "display": f"{mfg_n}종"},
        {"key": "accel", "label": "분양 가속 품목 수", "value": accel_n, "display": f"{accel_n}종"},
        {"key": "low_dist", "label": "장기 저분양 품목 수", "value": low_n, "display": f"{low_n}종"},
        {
            "key": "total_value",
            "label": "총 분양 환산금액",
            "value": total_value,
            "display": _fmt_money(total_value),
        },
        {
            "key": "reliability_ab",
            "label": "신뢰도 A·B 품목",
            "value": grade_ab,
            "display": f"{grade_ab}종",
        },
        {
            "key": "priced",
            "label": "가격 반영 품목",
            "value": valuation.get("priced_count", 0),
            "display": f"{valuation.get('priced_count', 0)}종",
        },
    ]

    summary_lines = [
        f"대상 품목 {managed}종을 기준으로 재고·분양 지표를 산출했습니다.",
        f"재고 없음(미보유) {zero_n}종은 소진 위험군에서 제외했습니다.",
        f"소진 예상: 1년 내 {deplete_1y}종 · 1~3년 {deplete_1_3y}종 · 3~5년 {deplete_3_5y}종 "
        f"(5년 이내 합계 {deplete_5y_total}종), 제조 우선검토 후보는 {mfg_n}종입니다.",
        f"분양 가속도(급가속·증가) 품목은 {accel_n}종, 장기 저분양/과다재고 후보는 {low_n}종입니다.",
        f"현 재고 기준 분양금액 환산 총액은 {_fmt_money(total_value)}입니다"
        + (
            f" (표준생약 {_fmt_money(by_type.get('표준생약'))}, "
            f"지표성분 {_fmt_money(by_type.get('지표성분'))}, "
            f"대조생약 {_fmt_money(by_type.get('대조생약'))})."
            if by_type
            else "."
        ),
    ]

    return {
        "kpis": kpis,
        "summary_lines": summary_lines,
        "valuation": valuation,
    }


def format_kpi_dashboard_markdown(dashboard: dict[str, Any]) -> str:
    """리포트 최상단용 KPI 대시보드 마크다운."""
    lines = [
        "## 1페이지 요약 대시보드 (핵심 KPI)",
        "",
        "| 지표 | 값 |",
        "| --- | --- |",
    ]
    for kpi in dashboard.get("kpis") or []:
        lines.append(f"| {kpi['label']} | {kpi['display']} |")
    lines.append("")
    lines.append("### 자동 종합 의견")
    for s in dashboard.get("summary_lines") or []:
        lines.append(f"- {s}")
    return "\n".join(lines)


def _monitoring_rows_to_markdown_table(rows: list[dict[str, Any]]) -> list[str]:
    """분양 가속 모니터링용 마크다운 표 (전수)."""
    headers = [
        "#",
        "한글명",
        "관리번호",
        "유형",
        "가속도",
        "비율",
        "연평균",
        "재고",
        "소진예상일시",
        "신뢰도",
    ]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for i, r in enumerate(rows, 1):
        ratio = r.get("acceleration_ratio")
        ratio_txt = f"{ratio:.2f}" if isinstance(ratio, (int, float)) else "-"
        ar = r.get("annual_rate")
        ar_txt = f"{ar:.2f}" if isinstance(ar, (int, float)) else "-"
        cells = [
            str(i),
            str(r.get("name_ko") or r.get("label") or "-").replace("|", "/"),
            str(r.get("manage_no") or "-").replace("|", "/"),
            str(r.get("std_type") or "-").replace("|", "/"),
            str(r.get("acceleration") or "-").replace("|", "/"),
            ratio_txt,
            ar_txt,
            format_qty_int(r.get("last_qty")),
            str(r.get("deplete_ym") or "-").replace("|", "/"),
            str(r.get("reliability") or "-").replace("|", "/"),
        ]
        lines.append("| " + " | ".join(cells) + " |")
    return lines


def format_accel_monitoring_markdown(monitoring: list[dict[str, Any]] | None) -> str:
    """표준 리포트 '분양 가속 모니터링' 섹션 — 항상 본문(표 또는 해당없음)을 포함."""
    rows = list(monitoring or [])
    surge_n = sum(1 for r in rows if r.get("acceleration") == "급가속")
    increase_n = sum(1 for r in rows if r.get("acceleration") == "증가")
    lines = [
        "## 분양 가속 모니터링",
        "",
        f"{ACCELERATION_FORMULA_KO}",
        "",
        f"급가속 {surge_n}건 · 증가 {increase_n}건 · 합계 {len(rows)}건 "
        "(마크다운 표 · 생략 없음)",
        "",
    ]
    if not rows:
        lines.append("해당 없음 (급가속·증가 품목 0건).")
        lines.append("")
        lines.extend(_monitoring_rows_to_markdown_table([]))
        return "\n".join(lines)
    lines.extend(_monitoring_rows_to_markdown_table(rows))
    return "\n".join(lines)


MANDATORY_REPORT_SECTION_ORDER = (
    "summary",
    "deplete",
    "missing",
    "manufacture",
    "accel",
    "compendium",
)


_CATALOG_SECTION_KEYS = frozenset(
    {"deplete", "missing", "manufacture", "accel", "compendium"}
)
_TRUNCATION_CELL_MARKERS = ("...", "…", "(중략)", "중략", "(생략)", "생략")


def _report_section_is_truncated(markdown: str) -> bool:
    """AI가 목록 중간을 줄인 경우(중략·생략 행)인지 판별."""
    text = markdown or ""
    if "중략" in text or "이하 생략" in text or "생략됨" in text:
        return True
    if re.search(r"등\s*\d+\s*건", text):
        return True
    for line in text.splitlines():
        s = line.strip()
        if not s.startswith("|") or _is_md_table_sep(s):
            continue
        cells = _parse_md_table_cells(s)
        if any((c or "").strip() in _TRUNCATION_CELL_MARKERS for c in cells):
            return True
    return False


def _report_section_has_item_rows(markdown: str) -> bool:
    """표 데이터 행 또는 번호 목록 품목이 있는지."""
    for line in (markdown or "").splitlines():
        s = line.strip()
        if re.match(r"^\d+\.\s+\S", s):
            return True
        if not s.startswith("|") or _is_md_table_sep(s):
            continue
        cells = _parse_md_table_cells(s)
        if not cells:
            continue
        if cells[0] in ("#", "번호", "한글명"):
            continue
        if cells[0].isdigit():
            return True
        if any(c in ("급가속", "증가", "표준생약", "지표성분") for c in cells):
            return True
    return False


def _report_section_body_is_empty(markdown: str) -> bool:
    """리포트 섹션에 헤딩 외 본문이 사실상 없는지(플레이스홀더·빈 표) 판별."""
    body = (markdown or "").strip()
    if not body:
        return True
    if "본문이 없습니다" in body:
        return True
    content_lines: list[str] = []
    for line in body.splitlines():
        s = line.strip()
        if not s or re.match(r"^#{1,4}\s+", s):
            continue
        content_lines.append(s)
    if not content_lines:
        return True
    if _report_section_has_item_rows(body):
        return False
    non_table = [
        ln for ln in content_lines if not ln.startswith("|") and not _is_md_table_sep(ln)
    ]
    if any(
        "해당 없음" in ln
        or re.search(r"0\s*건", ln)
        or ln == "없음"
        or ln == "(해당 없음)"
        for ln in non_table
    ):
        return False
    if non_table:
        only_meta = all(
            re.search(
                r"마크다운 표|÷|합계\s*\d+건|급가속\s*\d+건|증가\s*\d+건|건 ·|산출|우선순위점수",
                ln,
            )
            for ln in non_table
        )
        if not only_meta:
            return False
    return True


def _strip_auto_summary_opinion_blocks(md_text: str) -> str:
    """마크다운에서 '자동 종합 의견' 블록을 제거 (KPI 대시보드 주입 시 중복 방지)."""
    if not (md_text or "").strip():
        return ""
    out: list[str] = []
    skipping = False
    for line in md_text.splitlines():
        hm = re.match(r"^(#{2,4})\s+(.+)$", line.strip())
        if hm:
            title = re.sub(r"[*_`]", "", hm.group(2)).strip()
            if "자동 종합 의견" in title or title in ("종합 의견", "핵심 요약"):
                skipping = True
                continue
            skipping = False
            out.append(line)
            continue
        if not skipping:
            out.append(line)
    return "\n".join(out).strip()


def _ensure_manufacture_priority_formula(md_text: str) -> str:
    """제조 검토 섹션에 우선순위 산출 공식이 없으면 상단에 고정 삽입."""
    body = (md_text or "").strip()
    if not body:
        return body
    if PRIORITY_FORMULA_KO.splitlines()[0] in body or "제조우선순위점수 f" in body:
        return body
    lines = body.splitlines()
    if lines and lines[0].startswith("##"):
        return "\n".join(
            [
                lines[0],
                "",
                "### 제조 우선순위 점수 산정 공식",
                "",
                PRIORITY_FORMULA_KO,
                "",
                *lines[1:],
            ]
        )
    return "\n".join(
        [
            "### 제조 우선순위 점수 산정 공식",
            "",
            PRIORITY_FORMULA_KO,
            "",
            body,
        ]
    )


def _unique_missing_compendium_examples(
    missing_items: list[Any],
    *,
    limit: int = 8,
) -> list[str]:
    """공정서 미보유 대표 예시 — 생약명(한글명) 기준 고유."""
    seen: set[str] = set()
    examples: list[str] = []
    for row in missing_items:
        if isinstance(row, dict):
            name = row.get("name_ko") or row.get("name_en") or row.get("label")
            kind = row.get("pharmacopoeia_kind") or ""
        else:
            name = getattr(row, "name_ko", None) or str(row)
            kind = getattr(row, "pharmacopoeia_kind", "") or ""
        key = str(name or "").strip()
        if not key or key in seen:
            continue
        seen.add(key)
        examples.append(f"{key}" + (f"({kind})" if kind else ""))
        if len(examples) >= limit:
            break
    return examples


def _extract_raw_section(md_text: str, section_key: str) -> str | None:
    """원문 마크다운에서 지정 섹션 ##·### 블록을 추출. 없으면 None."""
    if not (md_text or "").strip():
        return None
    lines = md_text.splitlines()
    collecting = False
    buf: list[str] = []
    found = False
    for line in lines:
        hm = re.match(r"^(#{2,3})\s+(.+)$", line.strip())
        if hm:
            title = hm.group(2).strip()
            key = _canonical_section_key(title)
            if key == section_key:
                collecting = True
                found = True
                buf = [line]
                continue
            if collecting:
                break
            continue
        if collecting:
            buf.append(line)
    if not found:
        return None
    return "\n".join(buf).strip()


def _strip_section_blocks(md_text: str, section_key: str) -> str:
    """마크다운에서 지정 섹션 ##/### 블록을 제거."""
    if not (md_text or "").strip():
        return ""
    out: list[str] = []
    skipping = False
    for line in md_text.splitlines():
        hm = re.match(r"^(#{2,3})\s+(.+)$", line.strip())
        if hm:
            title = hm.group(2).strip()
            if _canonical_section_key(title) == section_key:
                skipping = True
                continue
            skipping = False
            out.append(line)
            continue
        if not skipping:
            out.append(line)
    return "\n".join(out).strip()


def _manufacture_rows_to_markdown_table(rows: list[dict[str, Any]]) -> list[str]:
    headers = [
        "#",
        "한글명",
        "관리번호",
        "유형",
        "재고",
        "우선순위",
        "소진예상일시",
        "소진구간",
        "위험등급",
    ]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for i, r in enumerate(rows, 1):
        score = r.get("priority_score")
        score_txt = f"{score:.3f}" if isinstance(score, (int, float)) else str(score or "-")
        cells = [
            str(i),
            str(r.get("name_ko") or r.get("label") or "-").replace("|", "/"),
            str(r.get("manage_no") or "-").replace("|", "/"),
            str(r.get("std_type") or "-").replace("|", "/"),
            format_qty_int(r.get("last_qty")),
            score_txt,
            str(r.get("deplete_ym") or "-").replace("|", "/"),
            str(r.get("depletion_category") or "-").replace("|", "/"),
            _display_risk_grade(r).replace("|", "/"),
        ]
        lines.append("| " + " | ".join(cells) + " |")
    return lines


def format_depletion_markdown(category_items: dict[str, list[dict[str, Any]]] | None) -> str:
    """표준 리포트 '소진 예상' 섹션 — 구간별 전수 표."""
    items = category_items or {}
    lines = [
        "## 소진 예상",
        "",
        "소진 예상 구간별 전수 목록 (마크다운 표 · 생략 없음)",
        "",
    ]
    total = 0
    for cat_name in DEPLETION_CATEGORY_ORDER:
        if cat_name == ZERO_STOCK_CATEGORY:
            continue
        rows = items.get(cat_name) or []
        total += len(rows)
        lines.append(f"### 소진구간 {cat_name} ({len(rows)}건)")
        lines.append("")
        if rows:
            lines.extend(_rows_to_markdown_table(rows))
        else:
            lines.append("(해당 없음)")
        lines.append("")
    if total == 0:
        lines.insert(3, "소진 예상 대상 0건 (재고 없음 품목은 미보유 섹션 참조).")
        lines.insert(4, "")
    return "\n".join(lines).strip()


def format_missing_markdown(
    zero_stock_rows: list[dict[str, Any]] | None,
    missing_compendium_rows: list[dict[str, Any]] | None,
) -> str:
    """표준 리포트 '미보유' 섹션 — 재고 0 + 공정서 미보유 전수 표."""
    zero_rows = list(zero_stock_rows or [])
    miss_rows = list(missing_compendium_rows or [])
    lines = [
        "## 미보유(재고 없음·공정서 미보유)",
        "",
        f"재고 없음 {len(zero_rows)}건 · 공정서 미보유 {len(miss_rows)}건 (마크다운 표 · 생략 없음)",
        "",
        f"### 재고 없음(미보유) ({len(zero_rows)}건)",
        "",
    ]
    if zero_rows:
        lines.extend(_rows_to_markdown_table(zero_rows))
    else:
        lines.append("(해당 없음)")
    lines.extend(
        [
            "",
            f"### 공정서 미보유 표준품 ({len(miss_rows)}건)",
            "",
        ]
    )
    if miss_rows:
        lines.extend(_missing_rows_to_markdown_table(miss_rows))
    else:
        lines.append("(해당 없음)")
    return "\n".join(lines)


def format_manufacture_review_markdown(manufacture: dict[str, list[dict[str, Any]]] | None) -> str:
    """표준 리포트 '차년도 제조 검토' 섹션 — 유형별 우선순위 상위 전수 표."""
    mfg = manufacture or {}
    std_rows = list(mfg.get("표준생약") or [])
    ind_rows = list(mfg.get("지표성분") or [])
    lines = [
        "## 차년도 제조 검토",
        "",
        "### 제조 우선순위 점수 산정 공식",
        "",
        PRIORITY_FORMULA_KO,
        "",
        f"표준생약 {len(std_rows)}건 · 지표성분 {len(ind_rows)}건 (마크다운 표 · 생략 없음)",
        "",
        f"### 표준생약 ({len(std_rows)}건)",
        "",
    ]
    if std_rows:
        lines.extend(_manufacture_rows_to_markdown_table(std_rows))
    else:
        lines.append("(해당 없음)")
    lines.extend(["", f"### 지표성분 ({len(ind_rows)}건)", ""])
    if ind_rows:
        lines.extend(_manufacture_rows_to_markdown_table(ind_rows))
    else:
        lines.append("(해당 없음)")
    return "\n".join(lines)


def build_mandatory_section_markdown(
    section_key: str,
    flags: dict[str, Any] | None = None,
    match_result: dict[str, Any] | None = None,
) -> str:
    """고정 리포트 탭용 섹션 본문을 ai_flags·공정서 매칭에서 생성."""
    f = flags or {}
    if section_key == "summary":
        dashboard = f.get("dashboard") or {}
        if dashboard.get("kpis"):
            return format_kpi_dashboard_markdown(dashboard)
        return "## 1페이지 요약 대시보드 (핵심 KPI)\n\n(요약 KPI 데이터 없음)"
    if section_key == "deplete":
        return format_depletion_markdown(f.get("depletion_category_items"))
    if section_key == "missing":
        cat = f.get("depletion_category_items") or {}
        zero_rows = cat.get(ZERO_STOCK_CATEGORY) or []
        miss_rows = f.get("missing_compendium_items") or []
        return format_missing_markdown(zero_rows, miss_rows)
    if section_key == "manufacture":
        return _ensure_manufacture_priority_formula(
            format_manufacture_review_markdown(f.get("manufacture_candidates"))
        )
    if section_key == "accel":
        return format_accel_monitoring_markdown(f.get("monitoring_targets"))
    if section_key == "compendium":
        return format_compendium_stats_markdown(match_result)
    return ""


def _mandatory_section_has_source_data(
    section_key: str,
    flags: dict[str, Any] | None,
    match_result: dict[str, Any] | None,
) -> bool:
    f = flags or {}
    if section_key == "summary":
        return bool((f.get("dashboard") or {}).get("kpis"))
    if section_key == "deplete":
        cat = f.get("depletion_category_items") or {}
        return any(len(cat.get(c) or []) for c in DEPLETION_CATEGORY_ORDER if c != ZERO_STOCK_CATEGORY)
    if section_key == "missing":
        cat = f.get("depletion_category_items") or {}
        return bool(cat.get(ZERO_STOCK_CATEGORY)) or bool(f.get("missing_compendium_items"))
    if section_key == "manufacture":
        mfg = f.get("manufacture_candidates") or {}
        return bool(mfg.get("표준생약")) or bool(mfg.get("지표성분"))
    if section_key == "accel":
        return bool(f.get("monitoring_targets"))
    if section_key == "compendium":
        return match_result is not None
    return False


def _section_needs_mandatory_inject(
    section_key: str,
    raw_section: str | None,
    flags: dict[str, Any] | None,
    match_result: dict[str, Any] | None,
) -> bool:
    """빈 본문·중략 표이거나, 목록 섹션은 정량 전수 표로 항상 교체."""
    raw = raw_section or ""
    if raw_section is None or _report_section_body_is_empty(raw):
        return True
    if _report_section_is_truncated(raw):
        return True
    if section_key in _CATALOG_SECTION_KEYS and _mandatory_section_has_source_data(
        section_key, flags, match_result
    ):
        return True
    if _mandatory_section_has_source_data(section_key, flags, match_result):
        if not _report_section_has_item_rows(raw):
            return True
    return False


def ensure_mandatory_report_sections(
    md_text: str,
    flags: dict[str, Any] | None = None,
    match_result: dict[str, Any] | None = None,
) -> str:
    """리포트 필수 섹션(요약·소진·미보유·제조검토·가속·공정서) 본문을 항상 채운다."""
    text = (md_text or "").strip()
    to_inject: list[tuple[str, str]] = []
    for key in MANDATORY_REPORT_SECTION_ORDER:
        built = build_mandatory_section_markdown(key, flags, match_result)
        raw = _extract_raw_section(text, key) if text else None
        if _section_needs_mandatory_inject(key, raw, flags, match_result):
            to_inject.append((key, built))

    if not to_inject:
        return text

    cleaned = text
    for key, _ in to_inject:
        cleaned = _strip_section_blocks(cleaned, key)
    if any(key == "summary" for key, _ in to_inject):
        cleaned = _strip_auto_summary_opinion_blocks(cleaned)

    summary_blocks = [body for key, body in to_inject if key == "summary"]
    tail_blocks = [
        _ensure_manufacture_priority_formula(body) if key == "manufacture" else body
        for key, body in to_inject
        if key != "summary"
    ]

    parts: list[str] = []
    if summary_blocks:
        parts.extend(summary_blocks)
    if cleaned:
        parts.append(cleaned.rstrip())
    if tail_blocks:
        parts.extend(tail_blocks)
    return "\n\n".join(parts).rstrip() + "\n"


def ensure_accel_monitoring_in_report(
    md_text: str,
    monitoring: list[dict[str, Any]] | None = None,
) -> str:
    """(호환) 가속 섹션만 보정 — ensure_mandatory_report_sections 사용 권장."""
    flags = {"monitoring_targets": list(monitoring or [])}
    return ensure_mandatory_report_sections(md_text, flags=flags)


# 하위 호환 별칭
_accel_section_has_item_rows = _report_section_has_item_rows
_accel_section_body_is_empty = _report_section_body_is_empty


def collect_ai_analysis_flags(items: list[StockItem]) -> dict[str, Any]:
    """AI 프롬프트와 동일한 기준으로 고갈·가속·우선순위·환산액 플래그를 산출."""
    deplete_codes: list[str] = []
    surge_codes: list[str] = []
    by_code: dict[str, dict[str, Any]] = {}
    stats_cache: dict[str, dict[str, Any]] = {}
    name_ko_map = build_name_ko_stock_map(items)
    manufacture = select_manufacture_candidates(items, stats_cache=stats_cache)
    monitoring = select_monitoring_targets(items, stats_cache=stats_cache)
    long_term_low = select_long_term_low_items(items, stats_cache=stats_cache)
    categories = group_by_depletion_category(
        items, stats_cache=stats_cache, name_ko_map=name_ko_map
    )
    category_items = group_by_depletion_category_items(
        items, stats_cache=stats_cache, name_ko_map=name_ko_map
    )
    risk_grade_items = group_by_risk_grade(items, stats_cache=stats_cache)
    valuation = compute_inventory_valuation(items)

    for it in items_for_ai_analysis(items):
        stats = get_depletion_stats(it, stats_cache)
        stock_zero = bool(stats.get("stock_zero")) or is_zero_stock(it)
        name_zero = is_name_level_zero_stock(it, name_ko_map)
        flag = {
            "label": it.label,
            "name_ko": it.name_ko,
            "name_en": _cell_str(getattr(it, "name_en", None)),
            "manage_no": it.manage_no,
            "speed": stats["speed"],
            "years_left": stats["years_left"],
            "deplete_within_2y": bool(stats["deplete_within_2y"]) and not stock_zero,
            "deplete_within_5y": bool(stats["deplete_within_5y"]) and not stock_zero,
            "recent_surge": bool(stats["recent_surge"]) and not stock_zero,
            "deplete_ym": stats["deplete_ym"],
            "depletion_category": (
                ZERO_STOCK_CATEGORY
                if name_zero
                else (
                    "5년 초과/안정"
                    if stock_zero and not name_zero
                    else stats["depletion_category"]
                )
            ),
            "risk_grade": (
                "재고없음"
                if name_zero
                else (
                    risk_grade_from_years(stats.get("years_left"))
                    if not stock_zero
                    else "안정"
                )
            ),
            "reliability": stats["reliability"],
            "priority_score": stats["priority_score"],
            "annual_rate": stats["annual_rate"],
            "rate_change_ratio": stats["rate_change_ratio"],
            "acceleration": stats["acceleration"],
            "acceleration_ratio": stats["acceleration_ratio"],
            "long_term_low": stats["long_term_low"] and not stock_zero,
            "stock_value": stats.get("stock_value"),
            "unit_price": stats.get("unit_price"),
            "last_qty": to_qty_int(it.last_qty if it.last_qty is not None else it.current_qty),
            "stock_zero": stock_zero,
            "name_level_zero": name_zero,
            "name_group_has_stock": name_ko_group_has_stock(it, name_ko_map),
            "increase_segments_excluded": stats.get("increase_segments_excluded", 0),
        }
        if it.manage_no:
            by_code[it.manage_no] = flag
        if flag["deplete_within_5y"]:
            deplete_codes.append(it.manage_no or it.label)
        if flag["acceleration"] in ("급가속", "증가") and not stock_zero:
            surge_codes.append(it.manage_no or it.label)

    flags = {
        "by_code": by_code,
        "deplete_codes": [c for c in deplete_codes if c],
        "surge_codes": [c for c in surge_codes if c],
        "deplete_labels": [
            by_code[c]["label"] for c in deplete_codes if c in by_code
        ],
        "surge_labels": [
            by_code[c]["label"] for c in surge_codes if c in by_code
        ],
        "manufacture_candidates": manufacture,
        "monitoring_targets": monitoring,
        "long_term_low_items": long_term_low,
        "depletion_categories": categories,
        "depletion_category_items": category_items,
        "risk_grade_items": risk_grade_items,
        "valuation": valuation,
        "name_ko_stock_map": name_ko_map,
        "_depletion_stats_cache": stats_cache,
    }
    flags["dashboard"] = build_kpi_dashboard(items, flags)
    flags.pop("_depletion_stats_cache", None)
    flags["chat_analysis_maps"] = build_chat_analysis_maps(items, flags)
    return flags


def build_chat_analysis_maps(
    items: list[StockItem],
    flags: dict[str, Any] | None = None,
    match_result: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """챗봇용 구조화 JSON — 한글명 재고맵·영문명 공정서맵·위험 TOP10·소진 요약."""
    flags = flags or {}
    name_map = flags.get("name_ko_stock_map") or build_name_ko_stock_map(items)
    name_ko_stock = [
        {
            "name_ko": s["name_ko"],
            "total_qty": to_qty_int(s["total_qty"]),
            "has_stock": bool(s["has_stock"]),
            "lot_count": int(s["lot_count"]),
            "manage_nos": list(s.get("manage_nos") or [])[:12],
        }
        for s in sorted(name_map.values(), key=lambda x: x["name_ko"])
    ]
    en_groups = build_name_en_inventory_groups(items)
    name_en_groups = []
    for key, group in en_groups.items():
        rep = next((g for g in group if not is_zero_stock(g)), group[0])
        name_en_groups.append(
            {
                "identity_key": key,
                "name_en": _cell_str(getattr(rep, "name_en", None)),
                "name_ko": rep.name_ko,
                "lot_count": len(group),
                "manage_nos": [g.manage_no for g in group if g.manage_no][:12],
                "total_qty": to_qty_int(
                    sum(float(g.current_qty or 0) for g in group)
                ),
                "has_stock": any(not is_zero_stock(g) for g in group),
            }
        )

    risk_rows: list[dict[str, Any]] = []
    for grade in ("위험", "경계", "주의"):
        for r in (flags.get("risk_grade_items") or {}).get(grade) or []:
            risk_rows.append({**r, "risk_grade": grade})
    risk_rows.sort(
        key=lambda r: (
            float(r["years_left"]) if isinstance(r.get("years_left"), (int, float)) else 999.0,
            -float(r.get("priority_score") or 0),
        )
    )
    risk_top10 = [
        {
            "name_ko": r.get("name_ko"),
            "manage_no": r.get("manage_no"),
            "risk_grade": r.get("risk_grade"),
            "years_left": r.get("years_left"),
            "deplete_ym": r.get("deplete_ym"),
            "priority_score": r.get("priority_score"),
            "last_qty": r.get("last_qty"),
        }
        for r in risk_rows[:10]
    ]

    cat = flags.get("depletion_category_items") or {}
    depletion_summary = {
        k: len(cat.get(k) or [])
        for k in DEPLETION_CATEGORY_ORDER
    }
    match = match_result or flags.get("compendium_match") or {}
    match_stats = match.get("stats") or {}
    en_match_raw = (
        match.get("name_en_match_map")
        or flags.get("name_en_match_map")
        or {}
    )
    name_en_match_map = [
        {"identity_key": k, **(v if isinstance(v, dict) else {"value": v})}
        for k, v in (en_match_raw.items() if isinstance(en_match_raw, dict) else [])
    ]
    return {
        "name_ko_stock_map": name_ko_stock,
        "name_en_inventory_groups": name_en_groups,
        "name_en_match_map": name_en_match_map,
        "risk_top10": risk_top10,
        "depletion_summary": depletion_summary,
        "compendium_match_stats": {
            "inventory_matched_groups": match_stats.get("inventory_matched"),
            "missing_count": match_stats.get("missing_count"),
            "compendium_total": match_stats.get("compendium_total"),
            "unique_en_groups": match_stats.get("unique_inventory_groups"),
            "inventory_matched_lots": match_stats.get("inventory_matched_lots"),
        },
        "as_of": date.today().isoformat(),
    }


def extract_mentioned_codes_from_report(report: str, known_codes: list[str]) -> list[str]:
    """AI 리포트 본문에 등장한 관리번호를 추출 (긴 코드 우선)."""
    if not report:
        return []
    ordered = sorted({c for c in known_codes if c}, key=len, reverse=True)
    found: list[str] = []
    for code in ordered:
        if code in report:
            found.append(code)
    return found


def _format_catalog_row_line(i: int, r: dict[str, Any]) -> str:
    """전수 목록 한 줄(번호·한글명·관리번호·세부 수치) — 표 미사용 시 폴백."""
    name = r.get("name_ko") or r.get("label") or "-"
    code = r.get("manage_no") or "-"
    yl = r.get("years_left")
    yl_txt = f"{yl:.2f}" if isinstance(yl, (int, float)) else "-"
    ar = r.get("annual_rate")
    ar_txt = f"{ar:.2f}" if isinstance(ar, (int, float)) else "-"
    ratio = r.get("acceleration_ratio")
    ratio_txt = f"{ratio:.2f}" if isinstance(ratio, (int, float)) else "-"
    score = r.get("priority_score")
    score_txt = f"{score:.3f}" if isinstance(score, (int, float)) else str(score or "-")
    rel = r.get("reliability_grade") or r.get("reliability") or "-"
    return (
        f"{i}. {name} | 관리번호:{code} | 유형:{r.get('std_type') or '-'} | "
        f"재고:{format_qty_int(r.get('last_qty'))} | 연평균:{ar_txt} | "
        f"가속도:{r.get('acceleration') or '-'}/{ratio_txt} | "
        f"소진년:{yl_txt} | {r.get('deplete_ym') or '-'} | "
        f"소진구간:{r.get('depletion_category') or '-'} | "
        f"위험등급:{r.get('risk_grade') or '-'} | "
        f"우선:{score_txt} | 신뢰:{rel}"
    )


def _rows_to_markdown_table(
    rows: list[dict[str, Any]],
    *,
    include_stock: bool = True,
) -> list[str]:
    """목록성 데이터를 마크다운 표로 전수 나열 (정수 재고·소진예상일시)."""
    headers = ["#", "한글명", "관리번호", "유형"]
    if include_stock:
        headers.append("재고")
    headers.extend(["소진예상일시", "소진구간", "위험등급"])
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for i, r in enumerate(rows, 1):
        cells = [
            str(i),
            str(r.get("name_ko") or r.get("label") or "-").replace("|", "/"),
            str(r.get("manage_no") or "-").replace("|", "/"),
            str(r.get("std_type") or "-").replace("|", "/"),
        ]
        if include_stock:
            cells.append(format_qty_int(r.get("last_qty")))
        cells.extend(
            [
                str(r.get("deplete_ym") or "-").replace("|", "/"),
                str(r.get("depletion_category") or "-").replace("|", "/"),
                _display_risk_grade(r).replace("|", "/"),
            ]
        )
        lines.append("| " + " | ".join(cells) + " |")
    return lines


def _missing_rows_to_markdown_table(rows: list[dict[str, Any]]) -> list[str]:
    headers = ["#", "한글명", "기원(한글)", "기원(영문)", "공정서"]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for i, r in enumerate(rows, 1):
        cells = [
            str(i),
            str(r.get("name_ko") or r.get("name_en") or "-").replace("|", "/"),
            str(r.get("origin_ko") or "-").replace("|", "/"),
            str(r.get("origin_en") or "-").replace("|", "/"),
            str(
                r.get("pharmacopoeia_kind") or r.get("pharmacopoeia") or "-"
            ).replace("|", "/"),
        ]
        lines.append("| " + " | ".join(cells) + " |")
    return lines


def _format_full_catalog_block(title: str, rows: list[dict[str, Any]]) -> list[str]:
    """카테고리 전수 목록 블록 — 마크다운 표(생략 없음)."""
    lines = [
        f"{title} — 전수 {len(rows)}건 (마크다운 표 · 생략 없음)",
        "",
    ]
    if not rows:
        lines.append("(해당 없음)")
        return lines
    lines.extend(_rows_to_markdown_table(rows))
    return lines


# AI 프롬프트 Token Diet — 전수 금지, 요약·TOP만
AI_PROMPT_TOP_DEPLETION = 20
AI_PROMPT_TOP_MISSING = 15
AI_PROMPT_TOP_N = 15  # 하위 호환


def _format_top_catalog_block(
    title: str,
    rows: list[dict[str, Any]],
    *,
    n: int = AI_PROMPT_TOP_N,
) -> list[str]:
    """(호환) 카탈로그 TOP N — 챗봇 스냅샷 등에서 사용."""
    total = len(rows)
    top = list(rows[: max(0, n)])
    lines = [
        f"{title} — 전체 {total}건 중 핵심 TOP {len(top)}건",
        "",
    ]
    if not top:
        lines.append("(해당 없음)")
        return lines
    lines.extend(_rows_to_markdown_table(top))
    if total > len(top):
        lines.append(f"(외 {total - len(top)}건 생략)")
    return lines


def _compact_item_summary_line(i: int, r: dict[str, Any]) -> str:
    name = r.get("name_ko") or r.get("label") or "-"
    code = r.get("manage_no") or "-"
    qty = r.get("last_qty")
    qty_txt = format_qty_int(qty) if qty is not None else "-"
    years = r.get("years_left")
    years_txt = f"{years:.1f}" if isinstance(years, (int, float)) else "-"
    return (
        f"{i}. {name} | {code} | 유형:{r.get('std_type') or '-'} | "
        f"재고:{qty_txt} | 등급:{r.get('risk_grade') or _display_risk_grade(r)} | "
        f"잔여:{years_txt}년 | 소진:{r.get('deplete_ym') or '-'} | "
        f"우선:{r.get('priority_score') if r.get('priority_score') is not None else '-'}"
    )


def build_ai_token_diet_block(flags: dict[str, Any], items: list[StockItem] | None = None) -> str:
    """AI 분석용 초경량 요약 — 통계 + 제조/소진 TOP20 + 미보유 15개만.

    478품목 전수·추이 시계열·카테고리 전수 표는 절대 포함하지 않는다.
    """
    maps = flags.get("chat_analysis_maps") or {}
    risk_items = flags.get("risk_grade_items") or {}
    category_items = flags.get("depletion_category_items") or {}
    manufacture = flags.get("manufacture_candidates") or {}
    miss_rows = list(flags.get("missing_compendium_items") or [])
    zero_rows = list(category_items.get(ZERO_STOCK_CATEGORY) or [])
    dashboard = flags.get("dashboard") or {}

    danger_n = len(risk_items.get("위험") or [])
    caution_n = len(risk_items.get("경계") or []) + len(risk_items.get("주의") or [])
    stable_n = len(risk_items.get("안정") or [])
    total_n = len(items) if items is not None else sum(
        len(v or []) for v in (category_items.values() if category_items else [])
    )
    if not total_n and dashboard.get("kpis"):
        for kpi in dashboard["kpis"]:
            if kpi.get("key") == "managed":
                total_n = int(kpi.get("value") or 0)
                break
    normal_n = max(0, total_n - danger_n - caution_n) if total_n else stable_n

    match_stats = maps.get("compendium_match_stats") or {}
    matched = match_stats.get("inventory_matched_groups")
    missing_c = match_stats.get("missing_count")
    if missing_c is None:
        missing_c = len(miss_rows)
    zero_n = len(zero_rows)

    dep_sum = maps.get("depletion_summary") or {
        k: len(category_items.get(k) or []) for k in DEPLETION_CATEGORY_ORDER
    }

    # 소진 임박 + 제조 우선순위 통합 TOP20
    scored: list[dict[str, Any]] = []
    seen: set[str] = set()
    for grade in ("위험", "경계", "주의"):
        for r in risk_items.get(grade) or []:
            key = str(r.get("manage_no") or r.get("label") or "")
            if not key or key in seen:
                continue
            seen.add(key)
            scored.append({**r, "risk_grade": r.get("risk_grade") or grade})
    for cat in ("1년 이내", "2년 이내", "3년 이내"):
        for r in category_items.get(cat) or []:
            key = str(r.get("manage_no") or r.get("label") or "")
            if not key or key in seen:
                continue
            seen.add(key)
            scored.append(dict(r))
    for bucket in (manufacture.get("표준생약") or []) + (manufacture.get("지표성분") or []):
        key = str(bucket.get("manage_no") or bucket.get("label") or "")
        if not key or key in seen:
            continue
        seen.add(key)
        scored.append(dict(bucket))

    def _sort_key(r: dict[str, Any]) -> tuple:
        yl = r.get("years_left")
        yl_v = float(yl) if isinstance(yl, (int, float)) else 999.0
        pr = float(r.get("priority_score") or 0)
        return (yl_v, -pr)

    scored.sort(key=_sort_key)
    top20 = scored[:AI_PROMPT_TOP_DEPLETION]

    missing_pool = miss_rows if miss_rows else zero_rows
    missing_top = missing_pool[:AI_PROMPT_TOP_MISSING]

    payload = {
        "stats": {
            "total_items": total_n,
            "normal_or_stable": normal_n,
            "caution_boundary": caution_n,
            "danger": danger_n,
            "zero_stock_missing_inventory": zero_n,
            "compendium_matched_groups": matched,
            "compendium_missing_count": missing_c,
            "depletion_counts": {k: int(dep_sum.get(k) or 0) for k in DEPLETION_CATEGORY_ORDER},
        },
        "manufacture_or_depletion_top20": [
            {
                "name_ko": r.get("name_ko") or r.get("label"),
                "manage_no": r.get("manage_no"),
                "std_type": r.get("std_type"),
                "last_qty": r.get("last_qty"),
                "risk_grade": r.get("risk_grade") or _display_risk_grade(r),
                "years_left": r.get("years_left"),
                "deplete_ym": r.get("deplete_ym"),
                "priority_score": r.get("priority_score"),
            }
            for r in top20
        ],
        "missing_sample_15": [
            {
                "name_ko": r.get("name_ko") or r.get("label") or r.get("name"),
                "manage_no": r.get("manage_no"),
                "compendium": r.get("compendium") or r.get("source") or r.get("std_type"),
            }
            for r in missing_top
        ],
    }
    try:
        body = json.dumps(payload, ensure_ascii=False, separators=(",", ":"), default=str)
    except (TypeError, ValueError):
        body = str(payload)
    # 안전 상한 (~수 천 토큰)
    if len(body) > 8000:
        body = body[:7950] + "...(truncated)"

    lines = [
        "[Token Diet 요약 JSON — 전수/추이 금지, 이 JSON만 근거]",
        "```json",
        body,
        "```",
        "",
        f"[제조 우선·소진 임박 TOP{len(top20)}]",
    ]
    if top20:
        lines.extend(_compact_item_summary_line(i, r) for i, r in enumerate(top20, 1))
    else:
        lines.append("(해당 없음)")
    lines.append("")
    lines.append(f"[미보유 대표 {len(missing_top)}건]")
    if missing_top:
        for i, r in enumerate(missing_top, 1):
            lines.append(
                f"{i}. {r.get('name_ko') or r.get('label') or r.get('name') or '-'} | "
                f"{r.get('manage_no') or '-'} | "
                f"{r.get('compendium') or r.get('source') or r.get('std_type') or '-'}"
            )
    else:
        lines.append("(해당 없음)")
    return "\n".join(lines)


def _build_ai_risk_summary_block(
    flags: dict[str, Any],
    *,
    n: int = AI_PROMPT_TOP_N,
) -> list[str]:
    """하위 호환 — Token Diet 블록을 줄 목록으로 반환."""
    return build_ai_token_diet_block(flags).splitlines()


def _format_candidate_lines(rows: list[dict[str, Any]], *, full: bool = False) -> str:
    if not rows:
        return "없음"
    # 제조검토 대상도 마크다운 표로 전수 나열
    table_rows = []
    for r in rows:
        table_rows.append(
            {
                "name_ko": r.get("name_ko") or r.get("label"),
                "manage_no": r.get("manage_no"),
                "std_type": r.get("std_type"),
                "last_qty": r.get("last_qty"),
                "deplete_ym": r.get("deplete_ym"),
                "depletion_category": r.get("depletion_category"),
                "risk_grade": r.get("risk_grade")
                or (
                    "재고없음"
                    if r.get("stock_zero")
                    else risk_grade_from_years(r.get("years_left"))
                ),
            }
        )
    header = (
        f"전수 {len(rows)}건 (마크다운 표)"
        if full
        else f"상위 {len(rows)}건 (마크다운 표)"
    )
    return "\n".join([header, ""] + _rows_to_markdown_table(table_rows))


_FULL_LIST_HINTS = (
    "전체",
    "모두",
    "전부",
    "전수",
    "전량",
    "다 보여",
    "다보여",
    "다 알려",
    "다알려",
    "리스트",
    "목록",
    "빠짐",
    "누락",
    "전부 알려",
    "다 출력",
)
_FULL_LIST_COUNT_RE = re.compile(r"\d+\s*건")


def detect_full_list_intent(question: str) -> dict[str, Any]:
    """전수 목록 요청 감지 — 카테고리별 플래그."""
    q = (question or "").strip()
    wants_full = any(h in q for h in _FULL_LIST_HINTS) or bool(_FULL_LIST_COUNT_RE.search(q))
    want_monitoring = any(k in q for k in ("모니터링", "급가속", "가속", "급증"))
    want_depletion = any(
        k in q
        for k in ("소진", "고갈", "1년", "2년", "3년", "4년", "5년", "10년", "15년", "구간", "미보유")
    )
    want_manufacture = any(k in q for k in ("제조", "후보", "우선검토", "우선 검토"))
    want_risk = (
        "위험등급" in q
        or any(k in q for k in ("위험군", "경계군", "주의군", "안정군"))
        or (
            any(k in q for k in ("위험", "경계", "주의"))
            and any(k in q for k in ("등급", "군", "품목", "목록", "전체", "모두"))
        )
    )
    want_missing = any(
        k in q
        for k in (
            "미보유",
            "부재",
            "없는 품목",
            "없는품목",
            "미등록",
            "공정서에만",
            "보유하지",
            "안 가진",
            "안가진",
        )
    )

    any_category = (
        want_monitoring or want_depletion or want_manufacture or want_risk or want_missing
    )
    # 카테고리 없이 "전체/모두"만 있으면 주요 카탈로그 전부 제공
    dump_all = wants_full and not any_category
    # 미보유는 키워드만으로도 전수 출력 의도로 간주 (「미보유 알려줘」)
    missing_full = want_missing and (
        wants_full or any(k in q for k in ("알려", "보여", "리스트", "목록", "뭐", "어느"))
    )
    return {
        "wants_full": wants_full or missing_full,
        "monitoring": wants_full and (want_monitoring or dump_all),
        "depletion": wants_full and (want_depletion or dump_all),
        "manufacture": wants_full and (want_manufacture or dump_all),
        "risk": wants_full and (want_risk or dump_all),
        "missing_compendium": missing_full or (wants_full and (want_missing or dump_all)),
        "dump_all": dump_all,
    }


def filter_missing_compendium_items(
    rows: list[dict[str, Any]],
    question: str,
) -> list[dict[str, Any]]:
    """미보유 목록 꼬리질문 필터 (KHP/KP/키워드)."""
    if not rows:
        return []
    q = (question or "").strip()
    filtered = list(rows)

    want_khp = any(k in q for k in ("KHP", "khp", "생약규격집", "약전외"))
    want_kp = False
    if not want_khp:
        if re.search(r"(?<![A-Za-z])KP(?![A-Za-z])", q) or "대한민국약전" in q:
            want_kp = True
        elif "약전" in q and "규격집" not in q and "약전외" not in q:
            want_kp = True

    if want_khp:
        filtered = [
            r
            for r in filtered
            if (r.get("pharmacopoeia_kind") or "") == "KHP"
            or "KHP" in str(r.get("pharmacopoeia_tag") or "").upper()
            or "생약규격집" in str(r.get("pharmacopoeia") or "")
        ]
    elif want_kp:
        filtered = [
            r
            for r in filtered
            if (r.get("pharmacopoeia_kind") or "") == "KP"
            or (
                "KP" in str(r.get("pharmacopoeia_tag") or "")
                and "KHP" not in str(r.get("pharmacopoeia_tag") or "").upper()
            )
        ]

    stop = {
        "그", "중", "중에", "만", "골라", "골라줘", "관련", "품목", "품목만", "있어", "있니",
        "미보유", "부재", "전체", "모두", "알려줘", "보여줘", "리스트", "목록",
        "공정서", "수재", "우선", "검토", "대상", "은", "는", "이", "가", "의",
        "없는", "없는품목", "다", "좀", "해줘", "생약규격집", "약전외", "대한민국약전",
        "약전", "kp", "khp", "KHP", "KP", "수재품목",
    }
    tokens = [
        t
        for t in re.split(r"[\s,./|?!~·\-_:;]+", q)
        if t and len(t) >= 2 and t not in stop and t.upper() not in ("KP", "KHP")
    ]
    # 조사·접미 제거 (품목만 → 품목)
    cleaned: list[str] = []
    for t in tokens:
        tt = re.sub(r"(만|은|는|이|가|을|를|의|도|만)$", "", t)
        if tt and tt not in stop and len(tt) >= 2:
            cleaned.append(tt)
    tokens = cleaned
    # 한글 실명 또는 의미 있는 영문 토큰만 이름 필터에 사용 (only/list 등과 공정서 구분 충돌 방지)
    ascii_noise = {
        "only", "just", "from", "with", "that", "this", "list", "show", "full",
        "all", "the", "and", "for", "item", "items", "please",
    }
    name_tokens = [t for t in tokens if re.search(r"[가-힣]", t)] + [
        t
        for t in tokens
        if t.isascii() and t.isalpha() and len(t) >= 3 and t.lower() not in ascii_noise
    ]
    if name_tokens:
        out: list[dict[str, Any]] = []
        for r in filtered:
            blob = (
                f"{r.get('name_ko') or ''}{r.get('name_en') or ''}"
                f"{r.get('origin_ko') or ''}{r.get('origin_en') or ''}"
            )
            norm_blob = _norm_key(blob)
            if any(
                (t in blob) or (_norm_key(t) and _norm_key(t) in norm_blob)
                for t in name_tokens
            ):
                out.append(r)
        filtered = out
    return filtered


def attach_compendium_match_to_flags(
    flags: dict[str, Any] | None,
    match_result: dict[str, Any] | None,
) -> dict[str, Any]:
    """ai_flags에 공정서 통계·미보유 전수를 붙여 챗봇 스냅샷과 동기화."""
    out = dict(flags or {})
    if not match_result:
        out.setdefault("compendium_stats", {})
        out.setdefault("missing_compendium_items", [])
        out.setdefault("name_en_match_map", {})
        return out
    out["compendium_match"] = match_result
    out["compendium_stats"] = match_result.get("stats") or {}
    out["name_en_match_map"] = match_result.get("name_en_match_map") or {}
    items = match_result.get("missing_items")
    if items is None and match_result.get("missing"):
        items = [
            _missing_compendium_row(e, i) if isinstance(e, CompendiumEntry) else e
            for i, e in enumerate(match_result["missing"], 1)
        ]
    out["missing_compendium_items"] = list(items or [])
    # 챗봇용 구조화 맵에 영문명 공정서 매칭 반영
    maps = out.get("chat_analysis_maps")
    if isinstance(maps, dict):
        refreshed = dict(maps)
        refreshed["name_en_match_map"] = [
            {"identity_key": k, **(v if isinstance(v, dict) else {"value": v})}
            for k, v in (out["name_en_match_map"] or {}).items()
        ]
        st = out["compendium_stats"] or {}
        refreshed["compendium_match_stats"] = {
            "inventory_matched_groups": st.get("inventory_matched"),
            "missing_count": st.get("missing_count"),
            "compendium_total": st.get("compendium_total"),
            "unique_en_groups": st.get("unique_inventory_groups"),
            "inventory_matched_lots": st.get("inventory_matched_lots"),
        }
        out["chat_analysis_maps"] = refreshed
    return out


def build_ai_prompt(
    items: list[StockItem],
    *,
    compendium_context: str | None = None,
    compendium_match_report: str | None = None,
    flags: dict[str, Any] | None = None,
) -> str:
    """AI 분석 프롬프트 — Token Diet (통계·TOP20·미보유15만, 전수 직렬화 금지)."""
    if flags is None:
        flags = collect_ai_analysis_flags(items)
    as_of = date.today()
    dashboard = flags.get("dashboard") or build_kpi_dashboard(items, flags)
    # compendium_* 원문 장문은 토큰 폭주 원인 — 통계는 diet JSON에만 반영
    _ = (compendium_context, compendium_match_report)

    lines = [
        "너는 대한민국 약전(KP)·생약규격집(KHP) 생약표준품 수석 데이터 분석가이다.",
        f"기준일: {as_of.isoformat()}",
        "",
        "[규칙]",
        "- 아래 Token Diet 요약만 근거로 분석하세요. 전수 목록·엑셀 원행·추이는 제공되지 않습니다.",
        "- 수치를 임의로 바꾸거나 없는 품목을 지어내지 마세요.",
        "- 서술: ①현황 수치 → ②원인 → ③권고. '제언' 섹션 금지.",
        "- 전수 표(소진·미보유·모니터링)는 앱이 리포트에 자동 주입하므로 모델은 요약 해석만 작성.",
        "",
        "[출력]",
        "1. 최상단 ## 1페이지 요약 대시보드 (아래 KPI 표를 그대로 포함, 종합의견 5줄 내외 1회)",
        "2. ## 소진 예상 / ## 미보유 / ## 차년도 제조 검토 / ## 분양 가속 모니터링 / "
        "## 공정서 DB 매칭 및 수재 현황 헤딩으로 핵심 이슈만 서술",
        f"3. 제조우선순위 공식 인용: {PRIORITY_FORMULA_KO}",
        "",
        "[1페이지 요약 대시보드(핵심 KPI)]",
        format_kpi_dashboard_markdown(dashboard),
        "",
        build_ai_token_diet_block(flags, items),
        "",
        "위 요약만으로 한국어 마크다운 분석 리포트를 작성하세요. "
        "전수 표·원본 행을 재구성하지 마세요.",
    ]
    prompt = "\n".join(lines)
    # 절대 상한 — 무료 티어 TPM 보호
    if len(prompt) > 12000:
        prompt = prompt[:11950] + "\n... (prompt truncated for token diet)"
    return prompt



def find_items_by_partial_query(
    items: list[StockItem],
    question: str,
) -> list[StockItem]:
    """질문 문자열로 관리번호·한글명·영문명 부분일치(contains) 품목 검색.

    예: "탄시논" → "탄시논 IIA" 매칭.
    """
    q = (question or "").strip()
    if not q or not items:
        return []
    q_lower = q.lower()
    tokens = [t for t in re.split(r"[\s,./?!~·\-_:;|()\[\]{}]+", q) if len(t) >= 2]

    results: list[StockItem] = []
    seen: set[str] = set()
    for it in items:
        key = it.manage_no or it.label
        if key in seen:
            continue
        code = (it.manage_no or "").strip()
        ko = (it.name_ko or "").strip()
        en = ""
        if not _is_empty(getattr(it, "name_en", None)):
            en = _cell_str(it.name_en).strip()
        en_l = en.lower()
        label = (it.label or "").strip()

        hit = False
        if code and code.lower() in q_lower:
            hit = True
        elif ko and (ko in q or any(t in ko for t in tokens)):
            hit = True
        elif en and (en_l in q_lower or any(t.lower() in en_l for t in tokens)):
            hit = True
        elif label and (any(t in label for t in tokens) or label.lower() in q_lower):
            hit = True
        if not hit and tokens:
            for t in tokens:
                tl = t.lower()
                if (ko and t in ko) or (en and tl in en_l) or (code and tl in code.lower()):
                    hit = True
                    break
        if hit:
            seen.add(key)
            results.append(it)
    return results


def serialize_flags_snapshot(flags: dict[str, Any] | None) -> str:
    """챗봇용 사전 산출 스냅샷 — 카테고리 전수 목록 포함(표준 리포트와 동일 수치)."""
    if not flags:
        return "[재고 조사 스냅샷] 없음 (플래그 미제공)"

    by_code = flags.get("by_code") or {}
    dashboard = flags.get("dashboard") or {}
    valuation = flags.get("valuation") or {}
    categories = flags.get("depletion_categories") or {}
    category_items = flags.get("depletion_category_items") or {}
    risk_grade_items = flags.get("risk_grade_items") or {}
    manufacture = flags.get("manufacture_candidates") or {}
    monitoring = flags.get("monitoring_targets") or []
    long_term_low = flags.get("long_term_low_items") or []
    surge_only = [r for r in monitoring if r.get("acceleration") == "급가속"]
    increase_only = [r for r in monitoring if r.get("acceleration") == "증가"]

    lines = [
        "[재고 조사 스냅샷 — 특정 시점의 재고 조사 데이터(사전 산출 결과). "
        "표준 분석 리포트와 동일 수치. 이 스냅샷·실시간 조회 수치만 사용하세요]",
        "[용어] '스냅샷(Snapshot)' = 특정 시점의 재고 조사 데이터(코드로 산출된 정량 결과).",
        "[전수 목록 규칙] 아래 마크다운 표는 요약·생략 없이 전부 수록되어 있습니다. "
        "사용자가 전체/모두/전수/N건·미보유·1년 이내 소진 목록을 요청하면 "
        "해당 표를 그대로 출력하세요. '등 N건'·일부만 나열·환각 추가를 금지합니다. "
        "재고량은 정수, 소진 시점은 'YYYY년 MM월'만 표기하세요 "
        "(예: 2027년 03월. '소진예상일시'·'기준' 문구 금지).",
        f"- by_code 품목 수: {len(by_code)}",
        f"- 5년 이내 소진 후보: {len(flags.get('deplete_codes') or [])}건",
        f"- 가속(급가속/증가) 후보: {len(flags.get('surge_codes') or [])}건",
        f"- 모니터링 대상 전수: {len(monitoring)}건 (급가속 {len(surge_only)} + 증가 {len(increase_only)})",
        format_kpi_dashboard_markdown(dashboard) if dashboard else "- KPI 대시보드: 없음",
        "",
        f"- 환산 총액: {_fmt_money(valuation.get('total_value'))}",
    ]
    for tname, tval in (valuation.get("by_type") or {}).items():
        lines.append(f"  · {tname}: {_fmt_money(tval)}")

    comp_stats = flags.get("compendium_stats") or {}
    missing_comp = flags.get("missing_compendium_items") or []
    if comp_stats or missing_comp:
        lines.append("")
        lines.append("[스냅샷: 공정서 DB 매칭 통계]")
        lines.append(
            f"- 총 수재:{comp_stats.get('compendium_total', 0)} · "
            f"보유매칭:{comp_stats.get('inventory_matched', 0)} · "
            f"자동보정:{comp_stats.get('auto_corrected', 0)} · "
            f"미보유:{comp_stats.get('missing_count', len(missing_comp))}"
        )
        lines.append(
            f"[스냅샷: 공정서 미보유 품목 전수] {len(missing_comp)}건 "
            "(마크다운 표 · 생략 없음)"
        )
        lines.extend(_missing_rows_to_markdown_table(list(missing_comp)))

    lines.append("")
    lines.append("[스냅샷: 소진 기간 카테고리 — 건수 요약]")
    for cat_name, labels in categories.items():
        detail = category_items.get(cat_name) or []
        n = len(detail) if detail else len(labels or [])
        lines.append(f"- {cat_name}: {n}건")

    lines.append("")
    if category_items:
        for cat_name, rows in category_items.items():
            lines.extend(
                _format_full_catalog_block(f"[스냅샷: 소진구간 {cat_name}]", rows)
            )
            lines.append("")
    else:
        lines.append("[스냅샷: 소진구간 전수] 미제공")
        lines.append("")

    lines.append("")
    if risk_grade_items:
        for grade in ("위험", "경계", "주의", "안정"):
            rows = risk_grade_items.get(grade) or []
            lines.extend(
                _format_full_catalog_block(f"[스냅샷: 위험등급 {grade}]", rows)
            )
            lines.append("")
    else:
        lines.append("[스냅샷: 위험등급 전수] 미제공")
        lines.append("")

    lines.append("[스냅샷: 차년도 제조검토 — 표준생약] (전수)")
    lines.append(_format_candidate_lines(manufacture.get("표준생약") or [], full=True))
    lines.append("[스냅샷: 차년도 제조검토 — 지표성분] (전수)")
    lines.append(_format_candidate_lines(manufacture.get("지표성분") or [], full=True))
    lines.append("")

    lines.extend(
        _format_full_catalog_block("[스냅샷: 모니터링 대상 전체]", list(monitoring))
    )
    lines.append("")
    lines.extend(_format_full_catalog_block("[스냅샷: 모니터링 — 급가속만]", surge_only))
    lines.append("")
    lines.extend(_format_full_catalog_block("[스냅샷: 모니터링 — 증가만]", increase_only))

    if long_term_low:
        lines.append("")
        lines.extend(
            _format_full_catalog_block(
                "[스냅샷: 장기 저분양]",
                [
                    {
                        "name_ko": r.get("label"),
                        "label": r.get("label"),
                        "manage_no": r.get("manage_no"),
                        "std_type": r.get("std_type"),
                        "annual_rate": r.get("annual_rate"),
                        "years_left": r.get("years_left"),
                        "reliability": r.get("reliability"),
                        "stock_value": r.get("stock_value"),
                    }
                    for r in long_term_low
                ],
            )
        )

    lines.append("")
    lines.append("[스냅샷: 품목별 핵심 지표 — by_code 요약(참고)]")
    for i, (code, st) in enumerate(by_code.items()):
        if i >= 120:
            lines.append(f"... (이하 {len(by_code) - 120}종은 by_code에 존재, 전수 목록 섹션 우선)")
            break
        yl = st.get("years_left")
        yl_txt = f"{yl:.2f}" if isinstance(yl, (int, float)) else "-"
        rel = st.get("reliability") or {}
        rel_g = rel.get("grade") if isinstance(rel, dict) else None
        lines.append(
            f"- {code} | {st.get('name_ko') or st.get('label')} | 속도:{st.get('speed')} | "
            f"가속:{st.get('acceleration')} | 소진년:{yl_txt} | "
            f"소진월:{st.get('deplete_ym')} | 위험:{st.get('risk_grade')} | "
            f"f:{st.get('priority_score')} | 환산:{_fmt_money(st.get('stock_value'))} | "
            f"신뢰:{rel_g}"
        )

    text = "\n".join(lines)
    # 전수 목록이 잘리면 안 되므로 상한을 넉넉히 둠
    if len(text) > 120_000:
        text = text[:119_960] + "\n... (스냅샷 용량 상한 — 전수 목록 앞부분이 우선)"
    return text


def _stats_from_flags(
    it: StockItem,
    by_code: dict[str, Any],
    *,
    allow_estimate: bool = True,
) -> dict[str, Any]:
    """flags by_code 스냅샷 우선, 없을 때만 estimate_depletion."""
    if it.manage_no and it.manage_no in by_code:
        return by_code[it.manage_no]
    if allow_estimate:
        return estimate_depletion(it)
    return {}


def query_live_inventory_context(
    items: list[StockItem],
    question: str,
    *,
    table_df: pd.DataFrame | None = None,
    max_detail_items: int = 35,
    flags: dict[str, Any] | None = None,
) -> str:
    """사용자 질문에 맞춰 재고 데이터를 필터링한 컨텍스트.

    flags가 주어지면 collect_ai_analysis_flags를 재호출하지 않고 스냅샷을 재사용한다.
    """
    if not items:
        return "[실시간 재고 조회] 업로드된 재고 데이터가 없습니다."

    q = (question or "").strip()
    q_lower = q.lower()
    reuse_flags = flags is not None
    if flags is None:
        flags = collect_ai_analysis_flags(items)
    dashboard = flags.get("dashboard") or build_kpi_dashboard(items, flags)
    valuation = flags.get("valuation") or compute_inventory_valuation(items)
    by_code = flags.get("by_code") or {}
    full_intent = detect_full_list_intent(q)

    matched_items = find_items_by_partial_query(items, q)
    hit_codes = sorted(
        {it.manage_no for it in matched_items if it.manage_no},
        key=len,
        reverse=True,
    )
    hit_names = sorted(
        {it.name_ko for it in matched_items if it.name_ko},
        key=len,
        reverse=True,
    )
    codes = {it.manage_no for it in items if it.manage_no}
    names = {it.name_ko for it in items if it.name_ko}
    for c in codes:
        if c and c.lower() in q_lower and c not in hit_codes:
            hit_codes.append(c)
    for n in names:
        if n and n in q and n not in hit_names:
            hit_names.append(n)
    specific_hits = bool(matched_items or hit_codes or hit_names)

    want_deplete = any(k in q for k in ("고갈", "소진", "제조", "우선", "5년", "2년"))
    want_accel = any(k in q for k in ("가속", "급증", "모니터링", "속도"))
    want_low = any(k in q for k in ("저분양", "과다", "하향"))
    want_value = any(k in q for k in ("환산", "가격", "금액", "자산", "가치"))
    want_reliab = any(k in q for k in ("신뢰", "데이터 부족", "등급"))

    selected: list[StockItem] = []
    seen: set[str] = set()

    def _add(it: StockItem) -> None:
        key = it.manage_no or it.label
        if key in seen:
            return
        seen.add(key)
        selected.append(it)

    for it in matched_items:
        _add(it)
    if not matched_items:
        for it in items:
            if it.manage_no in hit_codes or it.name_ko in hit_names:
                _add(it)

    matched_items = list(selected)

    if not specific_hits:
        if want_deplete or full_intent.get("depletion"):
            for code in flags.get("deplete_codes") or []:
                it = next((x for x in items if x.manage_no == code), None)
                if it:
                    _add(it)
        if want_accel or full_intent.get("monitoring"):
            for row in flags.get("monitoring_targets") or []:
                it = next((x for x in items if x.manage_no == row.get("manage_no")), None)
                if it:
                    _add(it)
        if want_low:
            for row in flags.get("long_term_low_items") or []:
                it = next((x for x in items if x.manage_no == row.get("manage_no")), None)
                if it:
                    _add(it)
        if full_intent.get("risk"):
            for rows in (flags.get("risk_grade_items") or {}).values():
                for row in rows:
                    it = next((x for x in items if x.manage_no == row.get("manage_no")), None)
                    if it:
                        _add(it)
        if full_intent.get("manufacture"):
            mfg = flags.get("manufacture_candidates") or {}
            for rows in (mfg.get("표준생약") or [], mfg.get("지표성분") or []):
                for row in rows:
                    it = next((x for x in items if x.manage_no == row.get("manage_no")), None)
                    if it:
                        _add(it)

    if not selected:
        def _score(it: StockItem) -> float:
            st = _stats_from_flags(it, by_code, allow_estimate=not reuse_flags)
            try:
                return float(st.get("priority_score") or 0)
            except (TypeError, ValueError):
                return 0.0

        scored = sorted(items_for_ai_analysis(items), key=_score, reverse=True)
        for it in scored[:max_detail_items]:
            _add(it)

    # 전수 목록 요청 시 상세 절단하지 않음
    if full_intent.get("wants_full") and (
        full_intent.get("monitoring")
        or full_intent.get("depletion")
        or full_intent.get("manufacture")
        or full_intent.get("risk")
        or full_intent.get("missing_compendium")
    ):
        pass  # keep full selected
    elif specific_hits and matched_items:
        extras = [it for it in selected if it not in matched_items]
        selected = matched_items + extras[: max(0, max_detail_items - len(matched_items))]
    else:
        selected = selected[:max_detail_items]

    lines = [
        "[실시간 재고 데이터 재검토 결과 — 초기 리포트보다 이 수치를 우선하세요]",
        "아래는 코드로 계산된 정량 팩트만입니다. 이 수치 외 추론하지 마세요.",
        (
            "- 수치 출처: 사전 산출 스냅샷(by_code) 재사용"
            if reuse_flags
            else "- 수치 출처: 실시간 collect_ai_analysis_flags"
        ),
        f"- 기준 시각: {date.today().isoformat()}",
        f"- 대상 품목 {len(items)}종 · 질문 매칭/선별 상세 {len(selected)}종",
        f"- 전수 목록 요청 감지: {full_intent}",
        format_kpi_dashboard_markdown(dashboard),
        "",
    ]

    missing_all = list(flags.get("missing_compendium_items") or [])
    # 미보유 꼬리질문/전수 — 스냅샷 정규화 목록을 즉시 필터
    if full_intent.get("missing_compendium") or any(
        k in q for k in ("미보유", "부재", "없는 품목", "없는품목")
    ):
        filtered_missing = filter_missing_compendium_items(missing_all, q)
        lines.append(
            "[공정서 미보유 출력 지시] 아래 목록만 근거로 답하세요. "
            "1~N 번호 또는 마크다운 표로 생략 없이 출력하고 '등 N건'을 쓰지 마세요. "
            f"필터 결과 {len(filtered_missing)}건 / 전체 미보유 {len(missing_all)}건."
        )
        lines.append(
            f"[실시간: 공정서 미보유 필터 결과] {len(filtered_missing)}건 (마크다운 표)"
        )
        lines.extend(_missing_rows_to_markdown_table(list(filtered_missing)))
        lines.append("")

    if full_intent.get("wants_full"):
        lines.append(
            "[전수 목록 출력 지시] 사용자가 전체/모두/전수/N건을 요청했습니다. "
            "아래 첨부 전수 목록을 번호 매김(1~N) 또는 마크다운 표로 빠짐없이 출력하고, "
            "'등 N건' 요약·임의 생략·없는 품목 추가를 금지하세요. "
            "출력 건수는 스냅샷 건수와 일치해야 합니다."
        )
        lines.append("")
        monitoring = flags.get("monitoring_targets") or []
        if full_intent.get("monitoring"):
            surge_only = [r for r in monitoring if r.get("acceleration") == "급가속"]
            lines.extend(
                _format_full_catalog_block("[실시간 전수: 모니터링 대상 전체]", list(monitoring))
            )
            lines.append("")
            lines.extend(
                _format_full_catalog_block("[실시간 전수: 급가속만]", surge_only)
            )
            lines.append("")
        if full_intent.get("depletion"):
            for cat_name, rows in (flags.get("depletion_category_items") or {}).items():
                lines.extend(
                    _format_full_catalog_block(f"[실시간 전수: 소진구간 {cat_name}]", rows)
                )
                lines.append("")
        if full_intent.get("risk"):
            for grade in ("위험", "경계", "주의", "안정"):
                rows = (flags.get("risk_grade_items") or {}).get(grade) or []
                lines.extend(
                    _format_full_catalog_block(f"[실시간 전수: 위험등급 {grade}]", rows)
                )
                lines.append("")
        if full_intent.get("manufacture"):
            mfg = flags.get("manufacture_candidates") or {}
            lines.append("[실시간 전수: 차년도 제조검토 — 표준생약]")
            lines.append(_format_candidate_lines(mfg.get("표준생약") or [], full=True))
            lines.append("[실시간 전수: 차년도 제조검토 — 지표성분]")
            lines.append(_format_candidate_lines(mfg.get("지표성분") or [], full=True))
            lines.append("")
        if full_intent.get("missing_compendium") and missing_all:
            lines.append(
                f"[실시간 전수: 공정서 미보유 전체] {len(missing_all)}건 "
                "(필터 없이 원본 전수 · 위 필터 결과와 구분)"
            )
            for i, r in enumerate(missing_all, 1):
                lines.append(
                    f"{i}. {r.get('name_ko') or r.get('name_en') or '-'} | "
                    f"공정서:{r.get('pharmacopoeia_kind') or '-'} | "
                    f"태그:{r.get('pharmacopoeia_tag') or '-'}"
                )
            lines.append("")

    if hit_codes or hit_names:
        lines.append(
            f"- 질문에서 식별된 키: 관리번호={', '.join(hit_codes) or '없음'} / "
            f"한글명={', '.join(hit_names) or '없음'}"
        )

    if want_value or specific_hits:
        lines.append(f"- 환산 총액: {_fmt_money(valuation.get('total_value'))}")
        for tname, tval in (valuation.get("by_type") or {}).items():
            lines.append(f"  · {tname}: {_fmt_money(tval)}")

    if want_reliab:
        grade_counts: dict[str, int] = defaultdict(int)
        for it in items:
            st = _stats_from_flags(it, by_code, allow_estimate=not reuse_flags)
            rel = st.get("reliability") or {}
            g = rel.get("grade", "?") if isinstance(rel, dict) else "?"
            grade_counts[str(g)] += 1
        lines.append(
            "- 신뢰도 등급 분포: "
            + ", ".join(f"{g}={n}" for g, n in sorted(grade_counts.items()))
        )

    if matched_items:
        lines.append("")
        lines.append("[질문 매칭 품목 — 정량 팩트 전체 (절단 없음)]")
        for it in matched_items:
            stats = _stats_from_flags(it, by_code, allow_estimate=True)
            if it.unit_price is None:
                it.unit_price = _extract_unit_price_from_meta(it.extra_meta)
            hist = ", ".join(
                f"{format_year(p.change_date)}={p.quantity:g}" for p in it.corrected_points
            )
            yl = stats.get("years_left")
            yl_txt = f"{yl:.4f}년" if isinstance(yl, (int, float)) else "산출불가"
            rel = stats.get("reliability") or {}
            lines.append(
                f"- {it.label} | manage_no={it.manage_no} | std_type={it.std_type} | "
                f"last_qty={it.last_qty} | unit_price={it.unit_price} | "
                f"stock_value={_fmt_money(it.stock_value)} | "
                f"speed={stats.get('speed')} | annual_rate={stats.get('annual_rate')} | "
                f"past_rate={stats.get('past_rate')} | recent_rate={stats.get('recent_rate')} | "
                f"acceleration={stats.get('acceleration')} | "
                f"acceleration_ratio={stats.get('acceleration_ratio')} | "
                f"years_left={yl_txt} | deplete_ym={stats.get('deplete_ym')} | "
                f"depletion_category={stats.get('depletion_category')} | "
                f"risk_grade={stats.get('risk_grade') or risk_grade_from_years(stats.get('years_left'))} | "
                f"deplete_within_2y={stats.get('deplete_within_2y')} | "
                f"deplete_within_5y={stats.get('deplete_within_5y')} | "
                f"priority_score={stats.get('priority_score')} | "
                f"stock_risk={stats.get('stock_risk')} | "
                f"speed_n={stats.get('speed_n')} | accel_n={stats.get('accel_n')} | "
                f"reliability={rel.get('label') if isinstance(rel, dict) else rel} | "
                f"reliability_grade={rel.get('grade') if isinstance(rel, dict) else None} | "
                f"increase_segments_excluded={stats.get('increase_segments_excluded', 0)} | "
                f"추이:[{hist}]"
            )

    lines.append("")
    lines.append("[선별 품목 실시간 지표]")
    for it in selected:
        if matched_items and it in matched_items:
            continue
        stats = _stats_from_flags(it, by_code, allow_estimate=not reuse_flags)
        if not stats:
            continue
        hist = ", ".join(
            f"{format_year(p.change_date)}={p.quantity:g}" for p in it.corrected_points
        )
        yl = stats.get("years_left")
        yl_txt = f"{yl:.2f}년" if isinstance(yl, (int, float)) else "산출불가"
        rel = stats.get("reliability") or {}
        lines.append(
            f"- {it.label} | {it.std_type} | 최종:{it.last_qty} | "
            f"속도:{stats.get('speed')} | 연평균:{stats.get('annual_rate')} | "
            f"가속도:{stats.get('acceleration')} | 소진:{yl_txt}({stats.get('deplete_ym')}) | "
            f"구간:{stats.get('depletion_category')} | 우선순위:{stats.get('priority_score')} | "
            f"환산:{_fmt_money(stats.get('stock_value'))} | "
            f"신뢰도:{(rel.get('label') if isinstance(rel, dict) else rel)} | "
            f"증가구간제외:{stats.get('increase_segments_excluded', 0)} | 추이:[{hist}]"
        )

    if table_df is not None and not table_df.empty and (hit_codes or hit_names or matched_items):
        lines.append("")
        lines.append("[통합 표(DataFrame) 매칭 행]")
        df = table_df
        mask = None
        if "관리번호" in df.columns and hit_codes:
            mask = df["관리번호"].astype(str).isin(hit_codes)
        match_names = hit_names or [it.name_ko for it in matched_items if it.name_ko]
        if "한글명" in df.columns and match_names:
            name_mask = df["한글명"].astype(str).isin(match_names)
            for n in match_names:
                name_mask = name_mask | df["한글명"].astype(str).str.contains(
                    re.escape(n), regex=True, na=False
                )
            mask = name_mask if mask is None else (mask | name_mask)
        if mask is not None:
            sub = df.loc[mask]
            prefer = [
                c for c in (
                    "관리번호", "한글명", "영문명", "표준품구분", "재고", "잔고",
                    "가격", "가격(원)", "분양여부",
                )
                if c in df.columns
            ]
            other = [c for c in df.columns if c not in prefer and not str(c).startswith("_")]
            cols = prefer + other
            for _, row in sub.iterrows():
                brief = []
                for col in cols:
                    v = row.get(col)
                    if _is_empty(v):
                        continue
                    brief.append(f"{col}={_cell_str(v)}")
                lines.append("- " + " | ".join(brief))

    text = "\n".join(lines)
    if full_intent.get("wants_full"):
        if len(text) > 120_000:
            text = text[:119_960] + "\n... (전수 목록 용량 상한)"
    elif len(text) > 24_000 and not matched_items:
        text = text[:23_960] + "\n... (이하 생략)"
    elif len(text) > 40_000:
        text = text[:39_960] + "\n... (이하 생략)"
    return text


def build_followup_prompt(
    base_report: str,
    user_question: str,
    items: list[StockItem] | None = None,
    *,
    compendium_context: str | None = None,
    table_df: pd.DataFrame | None = None,
    flags: dict[str, Any] | None = None,
) -> str:
    """초기 리포트 이후 추가 질문용 — 사전 산출 스냅샷과 실시간 조회를 주입."""
    full_intent = detect_full_list_intent(user_question or "")
    lines = [
        "너는 대한민국 약전(KP) 및 생약규격집(KHP) 생약표준품 관리 분야의 "
        "최고 수석 데이터 분석가이다.",
        "사용자의 후속 질문에는 초기 리포트 문구만 반복하지 말고, "
        "아래 [구조화 분석 맵 JSON]·[재고 조사 스냅샷]·[실시간 재고 데이터 재검토 결과]·공정서 DB를 "
        "우선 근거로 다각도 교차 분석하세요.",
        "",
        "[답변 논리 구조 — 필수]",
        "단순 단답을 지양하고, 질문 의도에 맞춰 "
        "① 현황 수치 → ② 원인 분석 → ③ 권고 액션 플랜 3단계로 심도 있게 답하세요.",
        "",
        "[용어] '스냅샷(Snapshot)'은 특정 시점의 재고 조사 데이터(코드로 산출된 정량 결과)를 의미합니다. "
        "답변에서 스냅샷이라고 말할 때 이 뜻을 명확히 하세요.",
        "[할루시네이션 금지] 제공된 JSON·스냅샷·실시간 수치·공정서 DB에 없는 내용을 단정하지 마세요. "
        "불확실하면 '확실하지 않음' 또는 '추측입니다'라고 밝히세요.",
        "[형식] '분석 전문가의 제언' 섹션은 작성하지 마세요.",
        "[공정서 DB] 규격/기준 참조에만 사용하고, 재고·분양 수치 산출에는 쓰지 마세요.",
        "[수치 동기화] 표준 분석 리포트 생성 시 산출된 스냅샷 수치만 사용하세요. "
        "임의로 다시 계산·추정하지 마세요.",
        "[우선순위] 초기 리포트 문구와 스냅샷/실시간 수치가 다르면 스냅샷·실시간 수치를 사용하세요.",
        "아래는 코드로 계산된 정량 팩트만입니다. 이 수치 외 추론하지 마세요.",
        "",
        "[전수 목록·표 출력 — 필수]",
        "- 사용자가 '전체/모두/전수/전량/리스트/목록/N건 모두' 등을 요청하면, "
        "초기 리포트의 '… 등 N건' 요약만 반복하지 마세요.",
        "- '공정서 수재 품목 중 미보유 표준품 목록/개수', '1년 이내 소진 예상 목록' 질문 시 "
        "스냅샷·실시간의 해당 마크다운 표를 정수 재고량·소진시점(YYYY년 MM월) 포함으로 누락 없이 출력하세요.",
        "- 접기/토글 없이 마크다운 표로 1~N 전수 출력. '등 N건'·임의 생략·없는 품목 추가 금지.",
        "- 출력 건수는 해당 섹션의 '전수 N건'과 반드시 일치해야 합니다.",
        "- 공정서 미보유/부재 꼬리질문(KHP만, 키워드 포함 등)은 "
        "스냅샷·실시간의 '공정서 미보유' 필터 결과만 사용하세요.",
        f"- 이번 질문 전수 요청 감지: {full_intent}",
        "",
    ]

    chat_json = format_chat_analysis_maps_json(flags)
    if chat_json:
        lines.append(chat_json)
        lines.append("")

    if flags:
        lines.append(serialize_flags_snapshot(flags))
        lines.append("")
    else:
        lines.append("[재고 조사 스냅샷] 미제공 — 실시간 조회 결과만 사용")
        lines.append("")

    if items:
        lines.append(
            query_live_inventory_context(
                items, user_question, table_df=table_df, flags=flags
            )
        )
        lines.append("")
    else:
        lines.append("[실시간 재고 데이터 재검토 결과] 재고 데이터 없음")
        lines.append("")

    if compendium_context and compendium_context.strip():
        ctx = compendium_context.strip()
        if len(ctx) > 8000:
            ctx = ctx[:7980] + "\n... (이하 생략)"
        lines.append(ctx)
        lines.append("")

    report = (base_report or "").strip()
    if report:
        # 전수 요청 시 초기 리포트(요약본) 비중을 줄여 환각·요약 반복을 막음
        limit = 1500 if full_intent.get("wants_full") else 3500
        if len(report) > limit:
            report = report[: limit - 20] + "\n... (초기 리포트 일부만 첨부 — 전수 목록은 스냅샷 우선)"
        lines.append("[초기 분석 리포트 — 보조 참고(요약일 수 있음. 전수 목록은 스냅샷·실시간 전수 섹션 우선)]")
        lines.append(report)
        lines.append("")

    lines.append(f"[사용자 질문]\n{user_question.strip()}")
    lines.append("")
    if full_intent.get("wants_full"):
        lines.append(
            "지금 질문은 전수 목록 요청입니다. 한국어 마크다운으로, "
            "해당 카테고리 전수 목록을 1~N 번호 목록(또는 표)으로 누락 없이 출력하세요. "
            "스냅샷/실시간 전수 섹션 외 추론·재계산·생략 금지."
        )
    else:
        lines.append(
            "질문에 대해 한국어 마크다운으로, "
            "① 현황 수치 → ② 원인 분석 → ③ 권고 액션 플랜 구조로 답하세요. "
            "필요한 경우 스냅샷·JSON·실시간 수치(연도·수량·소진시점·"
            "가속도·환산액·신뢰도 등급·우선순위점수 등)를 명시해 주세요. "
            "스냅샷/JSON에 있는 수치 외 추론·재계산은 하지 마세요."
        )
    return "\n".join(lines)


def _md_inline_to_html(text: str) -> str:
    """Escape + simple **bold** / `code` inline markdown."""
    esc = html.escape(text)
    esc = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", esc)
    esc = re.sub(r"`([^`]+)`", r"<code>\1</code>", esc)
    return esc


_MD_TABLE_SEP_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)


def _is_md_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and "|" in s[1:]


def _is_md_table_sep(line: str) -> bool:
    return bool(_MD_TABLE_SEP_RE.match(line.strip()))


def _parse_md_table_cells(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _md_table_rows_to_html(table_lines: list[str]) -> str:
    """연속된 마크다운 표 행을 HTML <table>로 변환."""
    if not table_lines:
        return ""
    rows: list[list[str]] = []
    header: list[str] | None = None
    for i, line in enumerate(table_lines):
        if _is_md_table_sep(line):
            if i == 1 and rows:
                header = rows.pop(0)
            continue
        cells = _parse_md_table_cells(line)
        if cells:
            rows.append(cells)

    parts = [
        '<table class="md-table" cellspacing="0" cellpadding="0" '
        'style="border-collapse:collapse;width:100%;max-width:560px;'
        'margin:8px 0 14px;font-size:10.5pt;">'
    ]
    if header:
        parts.append("<thead><tr>")
        for cell in header:
            parts.append(
                '<th style="border:1px solid #cbd5e1;background:#e8eef7;'
                'padding:7px 12px;text-align:left;font-weight:600;">'
                f"{_md_inline_to_html(cell)}</th>"
            )
        parts.append("</tr></thead>")
    parts.append("<tbody>")
    for ri, row in enumerate(rows):
        bg = "#ffffff" if ri % 2 == 0 else "#f8fafc"
        parts.append("<tr>")
        for ci, cell in enumerate(row):
            align = "right" if ci == len(row) - 1 and len(row) == 2 else "left"
            weight = "600" if ci == len(row) - 1 and len(row) == 2 else "400"
            parts.append(
                f'<td style="border:1px solid #cbd5e1;background:{bg};'
                f'padding:6px 12px;text-align:{align};font-weight:{weight};">'
                f"{_md_inline_to_html(cell)}</td>"
            )
        parts.append("</tr>")
    parts.append("</tbody></table>")
    return "".join(parts)


def _collapse_comma_items_html(
    body: str,
    preview_n: int,
    *,
    section_id: str,
    expanded: bool,
) -> str:
    """긴 쉼표 구분 품목 목록을 앵커 토글로 접기/펼치기.

    QTextBrowser는 HTML <details>를 지원하지 않으므로 #expand:id / #collapse:id 링크를 쓴다.
    """
    parts = [p.strip() for p in body.split(",") if p.strip()]
    if len(parts) <= preview_n:
        return _md_inline_to_html(body)
    n = len(parts)
    link_style = "color:#1e3a5f;text-decoration:underline;font-weight:600;"
    if expanded:
        full = _md_inline_to_html(", ".join(parts))
        return (
            f"{full} "
            f"<a href=\"#collapse:{section_id}\" style=\"{link_style}\">▼ 접기</a>"
        )
    preview = ", ".join(parts[:preview_n])
    return (
        f"{_md_inline_to_html(preview)}, … "
        f"<a href=\"#expand:{section_id}\" style=\"{link_style}\">"
        f"▶ 전체 {n}개 품목 펼쳐보기</a>"
    )


CANONICAL_REPORT_NAV = (
    # (key, short_label, title_keywords) — 앞쪽 키워드 우선
    ("summary", "요약", ("대시보드", "KPI", "요약")),
    ("deplete", "소진", ("소진 예상", "소진기간", "소진 기간", "소진구간", "소진")),
    ("missing", "미보유", ("미보유", "재고 없음", "재고없음", "부재")),
    ("manufacture", "검토", ("제조검토", "제조 검토", "우선검토", "제조")),
    ("accel", "가속", ("분양 가속 모니터링", "모니터링", "급가속", "분양 가속", "가속")),
    ("compendium", "공정서", ("공정서", "수재", "매칭")),
)


def split_markdown_report_sections(
    md_text: str,
    *,
    flags: dict[str, Any] | None = None,
    match_result: dict[str, Any] | None = None,
    monitoring: list[dict[str, Any]] | None = None,
) -> list[dict[str, str]]:
    """마크다운 리포트를 고정 사이드 버튼(요약·소진·미보유·검토 등) 기준으로 묶음.

    ## / ### 헤딩을 모두 수집해 키워드로 분류하고, 버튼용 섹션은 항상 반환한다.
    본문이 비어 있으면 ai_flags·공정서 매칭 정량 산출 표로 채운다.
    """
    if flags is None and monitoring is not None:
        flags = {"monitoring_targets": monitoring}
    text = (md_text or "").strip()
    buckets: dict[str, list[str]] = {key: [] for key, _, _ in CANONICAL_REPORT_NAV}
    other: list[str] = []

    if text:
        lines = text.splitlines()
        cur_key: str | None = None
        cur_lines: list[str] = []

        def _flush() -> None:
            nonlocal cur_key, cur_lines
            body = "\n".join(cur_lines).strip()
            if not body:
                cur_lines = []
                return
            if cur_key and cur_key in buckets:
                buckets[cur_key].append(body)
            else:
                other.append(body)
            cur_lines = []

        for line in lines:
            hm = re.match(r"^(#{2,3})\s+(.+)$", line.strip())
            if hm:
                _flush()
                title = hm.group(2).strip()
                cur_key = _canonical_section_key(title)
                cur_lines = [line]
            else:
                cur_lines.append(line)
        _flush()

    sections: list[dict[str, str]] = []
    for key, short, _ in CANONICAL_REPORT_NAV:
        parts = buckets.get(key) or []
        title = {
            "summary": "1페이지 요약 대시보드",
            "deplete": "소진 예상",
            "missing": "미보유(재고 없음·공정서 미보유)",
            "manufacture": "차년도 제조 검토",
            "accel": "분양 가속 모니터링",
            "compendium": "공정서 DB 매칭",
        }.get(key, short)
        built = build_mandatory_section_markdown(key, flags, match_result)
        if parts:
            body = "\n\n".join(parts).strip()
            if not body.lstrip().startswith("#"):
                body = f"## {title}\n\n{body}"
            if key == "summary":
                body = _strip_auto_summary_opinion_blocks(body)
            if _section_needs_mandatory_inject(key, body, flags, match_result):
                body = built
            elif key == "manufacture":
                body = _ensure_manufacture_priority_formula(body)
        else:
            body = built
        sections.append(
            {
                "id": key,
                "title": title,
                "short": short,
                "markdown": body,
            }
        )

    # [기타] 탭 제거 — 잔여 본문은 요약 탭에 병합
    if other:
        body = _strip_auto_summary_opinion_blocks("\n\n".join(other).strip())
        if body:
            for sec in sections:
                if sec.get("id") == "summary":
                    base = str(sec.get("markdown") or "").rstrip()
                    sec["markdown"] = (base + "\n\n" + body).strip() if base else (
                        body if body.lstrip().startswith("#") else f"## 요약\n\n{body}"
                    )
                    break
    return sections


def _canonical_section_key(title: str) -> str | None:
    """헤딩 제목 → 고정 네비 키 (미보유·검토가 공정서/소진에 먹히지 않도록 순서 고정)."""
    t = re.sub(r"\([^)]*\)", "", title or "").strip()
    t = re.sub(r"[#*_`]", "", t).strip()
    # 미보유·제조검토를 공정서/소진보다 먼저
    for key, _short, keywords in CANONICAL_REPORT_NAV:
        for kw in keywords:
            if kw in t:
                return key
    return None


def _report_section_short_label(title: str) -> str:
    """사이드 버튼용 짧은 라벨."""
    key = _canonical_section_key(title)
    if key:
        for k, short, _ in CANONICAL_REPORT_NAV:
            if k == key:
                return short
    t = re.sub(r"\([^)]*\)", "", title or "").strip()
    t = re.sub(r"[#*_`]", "", t).strip()
    compact = re.sub(r"\s+", "", t)
    return compact[:3] if compact else "기타"


def markdown_report_to_collapsible_html(
    md_text: str,
    preview_n: int = 8,
    expanded_ids: set[str] | frozenset[str] | None = None,
) -> str:
    """마크다운 리포트를 HTML로 변환. 목록은 접기 없이 전수 표시(표 우선)."""
    if not md_text:
        return ""
    out: list[str] = []
    in_ul = False
    lines = md_text.splitlines()
    i = 0

    def _close_ul() -> None:
        nonlocal in_ul
        if in_ul:
            out.append("</ul>")
            in_ul = False

    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            _close_ul()
            out.append("<br/>")
            i += 1
            continue

        # 마크다운 표 블록 → HTML table
        if _is_md_table_row(line):
            _close_ul()
            block: list[str] = []
            while i < len(lines) and _is_md_table_row(lines[i].rstrip()):
                block.append(lines[i].rstrip())
                i += 1
            out.append(_md_table_rows_to_html(block))
            continue

        hm = re.match(r"^(#{1,4})\s+(.*)$", line.strip())
        if hm:
            _close_ul()
            level = len(hm.group(1))
            out.append(f"<h{level}>{_md_inline_to_html(hm.group(2))}</h{level}>")
            i += 1
            continue

        lm = re.match(r"^(\s*[-*+]|\s*\d+\.)\s+(.*)$", line)
        if lm:
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{_md_inline_to_html(lm.group(2))}</li>")
            i += 1
            continue

        _close_ul()
        out.append(f"<p>{_md_inline_to_html(line)}</p>")
        i += 1

    _close_ul()
    style = (
        "<div style=\"font-family:'Malgun Gothic','Segoe UI',sans-serif;"
        "font-size:10.5pt;line-height:1.65;color:#1a2332;padding:4px;\">"
    )
    return style + "\n".join(out) + "</div>"



def category_counts(items: list[StockItem]) -> dict[str, int]:
    counts: dict[str, int] = defaultdict(int)
    for it in items:
        counts[it.std_type or "미분류"] += 1
    return dict(counts)


# ---------------------------------------------------------------------------
# 공정서(규격) 참조 DB — 재고 시계열 파싱 대상 아님, AI 프롬프트 컨텍스트 전용
# ---------------------------------------------------------------------------

COMPENDIUM_NAME_KEYS = (
    "생약명(한글)", "생약명", "한글명", "품목명", "명칭", "표준품명", "name", "Name",
)
COMPENDIUM_NAME_EN_KEYS = (
    "생약명(영어)", "생약명(영문)", "영문명", "영문", "영어명", "name_en", "Name_EN", "English",
)
COMPENDIUM_ORIGIN_KO_KEYS = ("기원(한글)", "기원", "기원식물", "원식물", "origin")
COMPENDIUM_ORIGIN_EN_KEYS = ("기원(영어)", "기원(영문)", "origin_en", "Origin_EN")
COMPENDIUM_PHARMACOPOEIA_KEYS = ("공정서", "수재공정서", "pharmacopoeia", "Pharmacopoeia", "KP", "KHP")
COMPENDIUM_CODE_KEYS = ("관리번호", "품목코드", "코드", "code", "Code")


@dataclass
class CompendiumEntry:
    name_ko: str = ""
    name_en: str = ""
    origin_ko: str = ""
    origin_en: str = ""
    pharmacopoeia: str = ""
    raw: dict[str, Any] = field(default_factory=dict)


def _norm_key(s: Any) -> str:
    """소문자·공백·괄호·특수문자 제거 정규화 키 (차집합·퍼지 매칭용)."""
    if _is_empty(s):
        return ""
    import unicodedata

    text = unicodedata.normalize("NFKC", str(s)).strip().lower()
    # 괄호·구두점·공백·언더스코어 등 비문자 제거
    text = re.sub(r"[\s\W_]+", "", text, flags=re.UNICODE)
    return text


def _pharmacopoeia_kind(pharm: str) -> str:
    """공정서 구분: KP / KHP / 기타."""
    tag = _pharmacopoeia_tag_from_text(pharm)
    if "KHP" in tag:
        return "KHP"
    if "KP" in tag and "KHP" not in tag:
        return "KP"
    if not pharm:
        return ""
    upper = str(pharm).upper()
    if "KHP" in upper or "생약규격집" in str(pharm) or "약전외" in str(pharm):
        return "KHP"
    if re.search(r"(?<![A-Z])KP(?![A-Z])", upper) or "대한민국약전" in str(pharm):
        return "KP"
    return str(pharm).strip()[:24]


def _missing_compendium_row(e: CompendiumEntry, index: int = 0) -> dict[str, Any]:
    kind = _pharmacopoeia_kind(e.pharmacopoeia)
    tag = _pharmacopoeia_tag_from_text(e.pharmacopoeia)
    return {
        "index": index,
        "name_ko": e.name_ko,
        "name_en": e.name_en,
        "origin_ko": e.origin_ko,
        "origin_en": e.origin_en,
        "pharmacopoeia": e.pharmacopoeia,
        "pharmacopoeia_kind": kind,
        "pharmacopoeia_tag": tag,
        "norm_ko": _norm_key(e.name_ko),
        "norm_en": _norm_key(e.name_en),
        "label": e.name_ko or e.name_en or f"entry-{index}",
    }


def _pick_column(columns: Sequence[str], keywords: Sequence[str]) -> Optional[str]:
    for key in keywords:
        for col in columns:
            if key.lower() in str(col).lower():
                return str(col)
    return None


def _cell_field(row: pd.Series, col: Optional[str]) -> str:
    if not col:
        return ""
    val = row.get(col)
    if _is_empty(val):
        return ""
    return _cell_str(val).strip()


def _parse_compendium_entries(df: pd.DataFrame) -> list[CompendiumEntry]:
    columns = [str(c) for c in df.columns]
    col_ko = _pick_column(columns, COMPENDIUM_NAME_KEYS)
    col_en = _pick_column(columns, COMPENDIUM_NAME_EN_KEYS)
    col_orig_ko = _pick_column(columns, COMPENDIUM_ORIGIN_KO_KEYS)
    col_orig_en = _pick_column(columns, COMPENDIUM_ORIGIN_EN_KEYS)
    col_pharm = _pick_column(columns, COMPENDIUM_PHARMACOPOEIA_KEYS)

    entries: list[CompendiumEntry] = []
    for _, row in df.iterrows():
        name_ko = _cell_field(row, col_ko)
        name_en = _cell_field(row, col_en)
        if not name_ko and not name_en:
            continue
        raw = {str(c): row.get(c) for c in columns if not str(c).startswith("_")}
        entries.append(
            CompendiumEntry(
                name_ko=name_ko,
                name_en=name_en,
                origin_ko=_cell_field(row, col_orig_ko),
                origin_en=_cell_field(row, col_orig_en),
                pharmacopoeia=_cell_field(row, col_pharm),
                raw=raw,
            )
        )
    return entries


def load_compendium_excel(path: PathLike) -> dict[str, Any]:
    """공정서/규격 DB 엑셀을 DataFrame + 구조화 entries로 로드 (시계열·소급보정 없음)."""
    path = str(path)
    frames: list[pd.DataFrame] = []
    with open_excel_file(path) as xls:
        for sheet in xls.sheet_names:
            part = pd.read_excel(xls, sheet_name=sheet)
            part.columns = [str(c).strip() for c in part.columns]
            if part.empty:
                continue
            part = part.copy()
            part.insert(0, "_시트", sheet)
            frames.append(part)
    if not frames:
        raise ValueError(f"공정서 DB 시트에 데이터가 없습니다: {Path(path).name}")

    df = pd.concat(frames, ignore_index=True, sort=False)
    entries = _parse_compendium_entries(df)
    return {
        "file_path": path,
        "file_name": Path(path).name,
        "dataframe": df,
        "row_count": int(len(df)),
        "columns": [str(c) for c in df.columns],
        "sheet_count": len(frames),
        "entries": entries,
        "parsed_count": len(entries),
    }


def _pharmacopoeia_tag_from_text(pharm: str) -> str:
    """공정서 문자열 → '[KP 수재]' / '[KHP(생약규격집) 수재]'."""
    if not pharm:
        return ""
    text = str(pharm).strip()
    upper = text.upper()
    # 약전외·생약규격집 = KHP (KP와 혼동 방지: '약전외'를 먼저 판정)
    is_khp = (
        "KHP" in upper
        or "생약규격집" in text
        or "약전외" in text
        or "한약(생약)규격" in text
        or "한약규격집" in text
    )
    is_kp = (
        bool(re.search(r"(?<![A-Z])KP(?![A-Z])", upper))
        or ("대한민국약전" in text and "약전외" not in text)
        or text in ("약전", "KP", "kp")
    )
    tags: list[str] = []
    if is_kp:
        tags.append("KP")
    if is_khp:
        tags.append("KHP(생약규격집)")
    if not tags:
        m = re.search(r"\b([A-Z]{2,5})\b", upper)
        if m and m.group(1) not in ("IIA", "IIB", "III"):
            tags.append(m.group(1))
    if not tags:
        return f"[{text} 수재]" if text else ""
    return " ".join(f"[{t} 수재]" for t in tags)


def match_compendium_inventory(
    entries: list[CompendiumEntry],
    items: list[StockItem],
) -> dict[str, Any]:
    """공정서 entries ↔ 재고 품목 매칭.

    재고 측은 영문명(없으면 한글명)이 같으면 1건으로 그룹화한 뒤 매칭한다.
    매칭 키: 한글명 exact/fuzzy, 영문명 fuzzy, 한글명↔기원 교차.
    """
    corrections: list[dict[str, Any]] = []
    by_manage_no: dict[str, str] = {}
    by_label: dict[str, str] = {}
    matched_entry_ids: set[int] = set()
    name_en_match_map: dict[str, dict[str, Any]] = {}

    by_exact_ko: dict[str, list[CompendiumEntry]] = defaultdict(list)
    by_norm_ko: dict[str, list[CompendiumEntry]] = defaultdict(list)
    by_norm_en: dict[str, list[CompendiumEntry]] = defaultdict(list)
    by_norm_origin_ko: dict[str, list[CompendiumEntry]] = defaultdict(list)
    by_norm_origin_en: dict[str, list[CompendiumEntry]] = defaultdict(list)

    for e in entries:
        if e.name_ko:
            by_exact_ko[e.name_ko.strip()].append(e)
            nk = _norm_key(e.name_ko)
            if nk:
                by_norm_ko[nk].append(e)
        if e.name_en:
            ne = _norm_key(e.name_en)
            if ne:
                by_norm_en[ne].append(e)
        if e.origin_ko:
            no = _norm_key(e.origin_ko)
            if no:
                by_norm_origin_ko[no].append(e)
        if e.origin_en:
            noe = _norm_key(e.origin_en)
            if noe:
                by_norm_origin_en[noe].append(e)

    def _pick_first(cands: list[CompendiumEntry]) -> Optional[CompendiumEntry]:
        for c in cands:
            if id(c) not in matched_entry_ids:
                return c
        return cands[0] if cands else None

    def _pick_by_origin(
        cands: list[CompendiumEntry], stock_ko: str
    ) -> Optional[CompendiumEntry]:
        nk = _norm_key(stock_ko)
        scored: list[tuple[int, CompendiumEntry]] = []
        for c in cands:
            if id(c) in matched_entry_ids and len(cands) > 1:
                continue
            score = 0
            ok = _norm_key(c.origin_ko)
            oe = _norm_key(c.origin_en)
            ck = _norm_key(c.name_ko)
            if nk and ck and nk == ck:
                score += 3
            if nk and ok and (nk in ok or ok in nk):
                score += 2
            if nk and oe and (nk in oe or oe in nk):
                score += 1
            scored.append((score, c))
        scored.sort(key=lambda x: x[0], reverse=True)
        if scored and scored[0][0] > 0:
            return scored[0][1]
        return _pick_first(cands)

    groups = build_name_en_inventory_groups(items)
    unique_inventory_groups = len(groups)

    for group_key, group_items in groups.items():
        rep = next((g for g in group_items if not is_zero_stock(g)), group_items[0])
        stock_ko = (rep.name_ko or "").strip()
        stock_en = _cell_str(getattr(rep, "name_en", None)).strip()
        hit: Optional[CompendiumEntry] = None
        match_type = ""

        if stock_ko and stock_ko in by_exact_ko:
            cands = by_exact_ko[stock_ko]
            hit = _pick_by_origin(cands, stock_ko) if len(cands) > 1 else _pick_first(cands)
            match_type = "exact_ko"
        if hit is None and stock_ko:
            nk = _norm_key(stock_ko)
            if nk and nk in by_norm_ko:
                cands = by_norm_ko[nk]
                hit = _pick_by_origin(cands, stock_ko) if len(cands) > 1 else _pick_first(cands)
                match_type = "fuzzy_ko"
        if hit is None and stock_en:
            ne = _norm_key(stock_en)
            if ne and ne in by_norm_en:
                hit = _pick_first(by_norm_en[ne])
                match_type = "fuzzy_en"
        if hit is None and stock_ko:
            nk = _norm_key(stock_ko)
            if nk and nk in by_norm_origin_ko:
                hit = _pick_first(by_norm_origin_ko[nk])
                match_type = "name_origin_ko"
            elif nk and nk in by_norm_origin_en:
                hit = _pick_first(by_norm_origin_en[nk])
                match_type = "name_origin_en"
            else:
                for ok, cands in by_norm_origin_ko.items():
                    if nk and (nk in ok or ok in nk) and len(nk) >= 2:
                        hit = _pick_first(cands)
                        match_type = "name_origin_ko"
                        break
                if hit is None:
                    for oe, cands in by_norm_origin_en.items():
                        if nk and (nk in oe or oe in nk) and len(nk) >= 2:
                            hit = _pick_first(cands)
                            match_type = "name_origin_en"
                            break

        if hit is None:
            continue

        matched_entry_ids.add(id(hit))
        tag = _pharmacopoeia_tag_from_text(hit.pharmacopoeia)
        short = ""
        if tag:
            inners = re.findall(r"\[(.+?)\s*수재\]", tag)
            short = inners[0].strip() if inners else ""
            if not short:
                found = re.findall(r"\[([^\]]+)", tag)
                short = found[0].strip() if found else (hit.pharmacopoeia or "")
            if short.upper() == "KHP" or (
                short.upper().startswith("KHP") and "생약규격집" not in short
            ):
                short = "KHP(생약규격집)"
        elif hit.pharmacopoeia:
            short = hit.pharmacopoeia.strip()
            if "생약규격집" in short or "약전외" in short or "KHP" in short.upper():
                short = "KHP(생약규격집)"

        name_en_match_map[group_key] = {
            "name_en": stock_en,
            "name_ko": stock_ko,
            "matched_name_ko": hit.name_ko,
            "matched_name_en": hit.name_en,
            "match_type": match_type,
            "pharmacopoeia": hit.pharmacopoeia,
            "lot_count": len(group_items),
        }

        for it in group_items:
            if it.manage_no and short:
                by_manage_no[it.manage_no] = short
            if tag:
                by_label[it.label] = tag
            corrections.append(
                {
                    "stock_label": it.label,
                    "stock_name_ko": (it.name_ko or "").strip() or stock_ko,
                    "stock_name_en": _cell_str(getattr(it, "name_en", None)).strip() or stock_en,
                    "matched_name_ko": hit.name_ko,
                    "matched_name_en": hit.name_en,
                    "matched_origin_ko": hit.origin_ko,
                    "matched_origin_en": hit.origin_en,
                    "pharmacopoeia": hit.pharmacopoeia,
                    "match_type": match_type,
                    "identity_key": group_key,
                    "group_lot_count": len(group_items),
                }
            )

    missing: list[CompendiumEntry] = [
        e for e in entries if id(e) not in matched_entry_ids
    ]
    missing_items = [
        _missing_compendium_row(e, i) for i, e in enumerate(missing, 1)
    ]
    auto_corrected = sum(
        1
        for c in name_en_match_map.values()
        if c.get("match_type") and c.get("match_type") != "exact_ko"
    )
    exact_matched = sum(
        1 for c in name_en_match_map.values() if c.get("match_type") == "exact_ko"
    )
    stats = {
        "compendium_total": len(entries),
        "inventory_matched": len(name_en_match_map),
        "inventory_matched_lots": len(corrections),
        "unique_inventory_groups": unique_inventory_groups,
        "entries_matched": len(matched_entry_ids),
        "exact_matched": exact_matched,
        "auto_corrected": auto_corrected,
        "missing_count": len(missing_items),
    }

    return {
        "corrections": corrections,
        "missing": missing,
        "missing_items": missing_items,
        "stats": stats,
        "by_manage_no": by_manage_no,
        "by_label": by_label,
        "correction_count": len(corrections),
        "name_en_match_map": name_en_match_map,
    }





def format_chat_analysis_maps_json(flags: dict[str, Any] | None) -> str:
    """챗봇용 초경량 구조화 JSON (샘플 최소화 — TPM 보호)."""
    if not flags:
        return ""
    maps = flags.get("chat_analysis_maps")
    if not isinstance(maps, dict) or not maps:
        return ""
    payload: dict[str, Any] = {
        "as_of": maps.get("as_of"),
        "depletion_summary": maps.get("depletion_summary") or {},
        "compendium_match_stats": maps.get("compendium_match_stats") or {},
        "risk_top10": maps.get("risk_top10") or [],
    }
    ko_map = list(maps.get("name_ko_stock_map") or [])
    payload["name_ko_stock_summary"] = {
        "total_names": len(ko_map),
        "with_stock": sum(1 for s in ko_map if s.get("has_stock")),
        "without_stock": sum(1 for s in ko_map if not s.get("has_stock")),
        "sample_with_stock": [s for s in ko_map if s.get("has_stock")][:10],
        "sample_without_stock": [s for s in ko_map if not s.get("has_stock")][:10],
    }
    en_match = list(maps.get("name_en_match_map") or [])
    payload["name_en_match_map_sample"] = en_match[:15]
    en_groups = list(maps.get("name_en_inventory_groups") or [])
    payload["name_en_group_summary"] = {
        "unique_groups": len(en_groups),
        "sample": en_groups[:10],
    }
    try:
        body = json.dumps(payload, ensure_ascii=False, separators=(",", ":"), default=str)
    except (TypeError, ValueError):
        body = str(payload)
    if len(body) > 6000:
        body = body[:5950] + "...(JSON truncated)"
    return (
        "[구조화 분석 맵 JSON — 요약만]\n"
        "아래 JSON만 교차 분석 근거로 사용하세요.\n"
        f"```json\n{body}\n```"
    )


def export_markdown_report_to_docx(md_text: str, path: str | Path) -> None:
    """AI 리포트 → Word(.docx) 경량 변환 (무한루프·복잡한 표 파서 없음).

    헤딩·글머리·일반 문단만 처리. 마크다운 표는 한 줄 문단으로 안전하게 기록.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Word 내보내기에 python-docx가 필요합니다. pip install python-docx"
        ) from exc

    text = (md_text or "").strip()
    if not text:
        raise ValueError("내보낼 리포트 본문이 비어 있습니다.")

    doc = Document()
    doc.add_heading("생약표준품 재고 분석 및 소진 예측 AI 리포트", level=1)

    # UI 프리징 방지: 과도한 행 수 상한
    lines = text.splitlines()
    max_lines = 8000
    if len(lines) > max_lines:
        lines = lines[:max_lines]
        lines.append("… (이하 생략 — 내보내기 안전 상한)")

    for line in lines:
        clean = line.strip()
        if not clean:
            continue
        # HTML 잔여 태그 제거(간단)
        if clean.startswith("<") and clean.endswith(">"):
            continue
        if clean.startswith("### "):
            doc.add_heading(clean[4:].strip(), level=3)
        elif clean.startswith("## "):
            doc.add_heading(clean[3:].strip(), level=2)
        elif clean.startswith("# "):
            doc.add_heading(clean[2:].strip(), level=1)
        elif clean.startswith(("- ", "* ", "• ")):
            doc.add_paragraph(clean[2:].strip(), style="List Bullet")
        else:
            # 굵게(**…**) 표시는 평문으로 (정규식 재귀 파싱 회피)
            plain = clean.replace("**", "")
            doc.add_paragraph(plain)

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def format_compendium_stats_markdown(match_result: dict[str, Any] | None) -> str:
    """표준 리포트용 공정서 매칭·수재 현황 섹션 (동적 집계)."""
    if not match_result:
        return (
            "## 공정서 DB 매칭 및 수재 현황\n\n"
            "- 공정서 DB가 등록되지 않아 집계할 수 없습니다.\n"
        )
    stats = match_result.get("stats") or {}
    missing_items = match_result.get("missing_items") or []
    if not stats and match_result.get("missing") is not None:
        # 구버전 결과 호환
        missing = match_result.get("missing") or []
        corrections = match_result.get("corrections") or []
        stats = {
            "compendium_total": len(missing) + len({id(c) for c in corrections}),
            "inventory_matched": len(corrections),
            "auto_corrected": sum(
                1 for c in corrections if c.get("match_type") != "exact_ko"
            ),
            "missing_count": len(missing),
        }
        if not missing_items and missing:
            missing_items = [
                _missing_compendium_row(e, i) if isinstance(e, CompendiumEntry)
                else e
                for i, e in enumerate(missing, 1)
            ]

    total = int(stats.get("compendium_total") or 0)
    held = int(stats.get("inventory_matched") or 0)
    auto_n = int(stats.get("auto_corrected") or 0)
    miss_n = int(stats.get("missing_count") or len(missing_items))
    examples = _unique_missing_compendium_examples(missing_items, limit=8)
    ex_txt = ", ".join(examples) if examples else "(해당 없음)"

    groups_n = int(stats.get("unique_inventory_groups") or held)
    lots_n = int(stats.get("inventory_matched_lots") or held)
    lines = [
        "## 공정서 DB 매칭 및 수재 현황",
        "",
        f"- 총 공정서 수재 품목 수: **{total}건**",
        f"- 재고 엑셀 보유 매칭 품목 수(영문명 기준 품목군): **{held}건** "
        f"(고유 품목군 {groups_n}건 / 매칭 로트 {lots_n}건)",
        f"- 기원(한글/영문) 기반 자동 보정 매칭 품목 수: **{auto_n}건**",
        f"- 공정서 수재 품목 중 미보유(부재) 품목 총 건수: **{miss_n}건**",
        f"- 미보유 대표 예시: {ex_txt}",
        "",
        f"### 공정서 미보유 표준품 전수 ({miss_n}건)",
        "",
    ]
    if missing_items:
        dict_rows = [
            r if isinstance(r, dict) else _missing_compendium_row(r, i)
            for i, r in enumerate(missing_items, 1)
        ]
        lines.extend(_missing_rows_to_markdown_table(dict_rows))
    else:
        lines.append("(해당 없음)")
    lines.append("")
    return "\n".join(lines)


def format_compendium_match_report(match_result: dict[str, Any]) -> str:
    """보정 매칭·미보유 공정서 표준품 보고 텍스트 (표준 리포트 정식 섹션 포함)."""
    corrections = match_result.get("corrections") or []
    missing_items = match_result.get("missing_items") or []
    if not missing_items and match_result.get("missing"):
        missing_items = [
            _missing_compendium_row(e, i) if isinstance(e, CompendiumEntry) else e
            for i, e in enumerate(match_result["missing"], 1)
        ]

    lines = [
        format_compendium_stats_markdown(match_result).rstrip(),
        "",
        "[표준 리포트 작성 지시] 위 '공정서 DB 매칭 및 수재 현황'과 미보유 전수 마크다운 표를 "
        "리포트 본문에 수치 변경 없이 그대로 포함하세요. 접기/토글 없이 표로 전수 나열합니다.",
        "",
        f"[공정서 매칭 보정 상세] 총 {len(corrections)}건",
    ]
    if corrections:
        for i, c in enumerate(corrections[:30], 1):
            lines.append(
                f"{i}. {c.get('stock_label')} ← {c.get('matched_name_ko') or c.get('matched_name_en')} "
                f"(type={c.get('match_type')}) | 기원:{c.get('matched_origin_ko') or '-'} | "
                f"공정서:{c.get('pharmacopoeia') or '-'}"
            )
        if len(corrections) > 30:
            lines.append(f"- … 보정 상세 외 {len(corrections) - 30}건")
    else:
        lines.append("- 보정 매칭 없음")

    return "\n".join(lines)


def lookup_pharmacopoeia_tag(
    item: StockItem,
    match_result: dict[str, Any] | None,
) -> str:
    """품목에 대한 '[KP 수재]' 형태 태그 또는 빈 문자열."""
    if not match_result:
        return ""
    by_label = match_result.get("by_label") or {}
    if item.label in by_label:
        return by_label[item.label]
    by_manage = match_result.get("by_manage_no") or {}
    code = by_manage.get(item.manage_no or "")
    if code:
        return _pharmacopoeia_tag_from_text(str(code))
    return ""


def _row_to_brief(row: pd.Series, columns: list[str], max_cols: int = 12) -> str:
    parts: list[str] = []
    for col in columns[:max_cols]:
        if str(col).startswith("_"):
            continue
        val = row.get(col)
        if _is_empty(val):
            continue
        text = _cell_str(val)
        if len(text) > 80:
            text = text[:77] + "..."
        parts.append(f"{col}={text}")
    return " | ".join(parts)


def format_compendium_context(
    df: pd.DataFrame | None,
    items: list[StockItem] | None = None,
    *,
    max_rows: int = 60,
    max_chars: int = 12_000,
    meta: dict[str, Any] | None = None,
    match_result: dict[str, Any] | None = None,
) -> str:
    """AI 프롬프트용 공정서 DB 요약·매칭 행 텍스트.

    main.py 에서 df(+meta)만 넘겨도 동작. items가 있으면 구조화 매칭을 우선 사용.
    """
    if df is None or df.empty:
        return ""

    columns = [str(c) for c in df.columns]
    entries: list[CompendiumEntry] = []
    if meta and meta.get("entries"):
        entries = list(meta["entries"])
    else:
        entries = _parse_compendium_entries(df)

    lines = [
        "[공정서/규격 참조 DB — 재고 시계열이 아님. 규격·기준 정보 참조 전용]",
        f"- 파일: {(meta or {}).get('file_name', '')}",
        f"- 행 수: {len(df)} · 구조화 entries: {len(entries)} · 열: {', '.join(columns[:20])}"
        + (" ..." if len(columns) > 20 else ""),
        "- 사용 규칙: 재고/분양 수치 산출에 쓰지 말고, 품목 규격·기준 설명에만 인용하세요.",
        "- DB에 없는 규격 정보는 '공정서 DB에 해당 항목 없음'으로 명시하세요.",
    ]

    if entries:
        lines.append(
            f"- entries 요약(상위 {min(8, len(entries))}): "
            + ", ".join(
                (e.name_ko or e.name_en) + (f"/{e.pharmacopoeia}" if e.pharmacopoeia else "")
                for e in entries[:8]
            )
        )

    local_match = match_result
    if local_match is None and items and entries:
        local_match = match_compendium_inventory(entries, items)

    if local_match and (local_match.get("corrections") or local_match.get("missing") is not None):
        lines.append("")
        report = format_compendium_match_report(local_match)
        # 컨텍스트 길이 보호를 위해 리포트도 포함하되 아래에서 max_chars로 절단
        lines.append(report)

    name_col = _pick_column(columns, COMPENDIUM_NAME_KEYS)
    code_col = _pick_column(columns, COMPENDIUM_CODE_KEYS)
    keys_name = {it.name_ko.strip() for it in (items or []) if it.name_ko}
    keys_code = {it.manage_no.strip() for it in (items or []) if it.manage_no}

    def _row_matches(row: pd.Series) -> bool:
        if name_col and not _is_empty(row.get(name_col)):
            nm = _cell_str(row.get(name_col))
            if nm in keys_name or any(nm and (nm in k or k in nm) for k in keys_name):
                return True
        if code_col and not _is_empty(row.get(code_col)):
            cd = _cell_str(row.get(code_col))
            if cd in keys_code:
                return True
        return False

    matched_rows: list[pd.Series] = []
    if keys_name or keys_code:
        for _, row in df.iterrows():
            if _row_matches(row):
                matched_rows.append(row)

    lines.append("")
    if matched_rows:
        lines.append(
            f"[재고 품목과 매칭된 공정서 행] ({min(len(matched_rows), max_rows)}/{len(matched_rows)}건 표시)"
        )
        for row in matched_rows[:max_rows]:
            lines.append(f"- {_row_to_brief(row, columns)}")
    elif not local_match:
        lines.append("[공정서 DB 미리보기] (재고 품목과 자동 매칭된 행 없음 — 상위 행 샘플)")
        preview_n = min(15, len(df), max_rows)
        for _, row in df.head(preview_n).iterrows():
            lines.append(f"- {_row_to_brief(row, columns)}")

    text = "\n".join(lines)
    if len(text) > max_chars:
        text = text[: max_chars - 20] + "\n... (이하 생략)"
    return text


# 3D 표시용 예상 소진기간 상한 (AI 추이와 동일 산출, 초장기는 상한에 묶음)
YEARS_LEFT_DISPLAY_CAP = 50.0


def scatter_cat_label(std_type: str, std_type_raw: Any = None) -> str:
    """3D 범례용 표준품구분 (접두어 없이 표준생약/지표성분/대조품)."""
    text = std_type or _cell_str(std_type_raw)
    if not text:
        return "미분류"
    for key, label in (
        ("표준생약", "표준생약"),
        ("지표성분", "지표성분"),
        ("대조품", "대조품"),
        ("대조생약", "대조품"),
    ):
        if key in text:
            return label
    match = re.search(r"\(([^)]+)\)", text)
    if match:
        return match.group(1).strip() or text
    return text


def build_scatter3d_record(
    item: StockItem,
    ai_flags: dict[str, Any] | None = None,
    mentioned_codes: set[str] | None = None,
) -> Optional[dict[str, Any]]:
    """AI 추이(estimate_depletion) 기준 3D 레코드. 감소 추이가 없으면 None."""
    if not item.has_stock_change:
        return None

    pts = item.corrected_points
    if len(pts) < 2 or item.first_qty is None or item.last_qty is None:
        return None

    flag = (ai_flags or {}).get(item.manage_no) if ai_flags else None
    if flag is not None:
        deplete = bool(flag.get("deplete_within_5y"))
        surge = bool(flag.get("recent_surge"))
        speed = flag.get("speed") or ""
        years_left = flag.get("years_left")
        annual_rate = flag.get("annual_rate")
    else:
        stats = estimate_depletion(item)
        deplete = bool(stats.get("deplete_within_5y"))
        surge = bool(stats.get("recent_surge"))
        speed = stats.get("speed")
        years_left = stats.get("years_left")
        annual_rate = stats.get("annual_rate")

    # 감소 없음/증가만 있는 품목은 천년대 runway로 왜곡되므로 제외
    if annual_rate is None or annual_rate <= 1e-9 or years_left is None:
        return None

    years_raw = float(years_left)
    years_plot = min(years_raw, YEARS_LEFT_DISPLAY_CAP)
    capped = years_raw > YEARS_LEFT_DISPLAY_CAP

    qtys = [p.quantity for p in pts]
    total_variation = sum(abs(qtys[i + 1] - qtys[i]) for i in range(len(qtys) - 1))
    net_drop = float(item.first_qty - item.last_qty)
    init_qty = float(item.first_qty)
    balance = float(item.last_qty)
    days = (pts[-1].change_date - pts[0].change_date).days or 1
    elapsed_years = max(days / 365.25, 1 / 365.25)
    if abs(init_qty) > 1e-12:
        decrease_rate = net_drop / init_qty * 100.0
    else:
        decrease_rate = 0.0

    mentioned = bool(mentioned_codes and item.manage_no in mentioned_codes)

    return {
        "cat": scatter_cat_label(item.std_type, item.std_type_raw),
        "code": item.manage_no,
        "name": item.name_ko,
        "balance": round(balance, 6),
        "initQty": round(init_qty, 6),
        "netDrop": round(net_drop, 6),
        "totalVariation": round(float(total_variation), 6),
        "decreaseRate": round(float(decrease_rate), 4),
        "annualRate": round(float(annual_rate), 6),
        "runway": round(years_plot, 6),
        "elapsedYears": round(float(elapsed_years), 6),
        "depleteWithin5y": deplete,
        "recentSurge": surge,
        "speed": speed or "",
        "yearsLeft": round(years_plot, 4),
        "yearsLeftRaw": round(years_raw, 4),
        "yearsCapped": capped,
        "aiMentioned": mentioned,
    }


def build_scatter3d_records(
    items: list[StockItem],
    ai_flags: dict[str, Any] | None = None,
    mentioned_codes: set[str] | list[str] | None = None,
) -> list[dict[str, Any]]:
    """유효한 3D 산점도 레코드만 모은다."""
    flag_map = None
    if isinstance(ai_flags, dict) and "by_code" in ai_flags:
        flag_map = ai_flags["by_code"]
    elif isinstance(ai_flags, dict):
        flag_map = ai_flags
    mentioned = set(mentioned_codes or [])
    records: list[dict[str, Any]] = []
    for item in items:
        rec = build_scatter3d_record(item, ai_flags=flag_map, mentioned_codes=mentioned)
        if rec is not None:
            records.append(rec)
    return records
